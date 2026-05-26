"""
Análisis Financiero Proyección 2027 UnionX — VERSIÓN 3
Correcciones aplicadas v3:
- Control Gestión NO es gerencia (sub-gerencia Finanzas)
- Gabriela se mantiene (no sale)
- Costo/pedido actual real $2.812 (no allocated $2.887)
- Matriz extendida 7 escenarios MC 27%→33% con apertura completa de variables
- Análisis 2026 cambio operador completo (3 sub-escenarios)
- 2027 por escenario con costo/ped in-house vs operador
- Punto equilibrio largo plazo con 3 casos
"""
import os, sys, warnings
warnings.filterwarnings('ignore')
if sys.stdout.encoding != 'utf-8':
    try: sys.stdout.reconfigure(encoding='utf-8')
    except: pass
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import openpyxl
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"G:\Mi unidad\TRABAJO\RESPALDO\OPERACIONES\UNION X - IA"
LOGO = os.path.join(BASE, "data", "branding", "unionx_logo.png")
PLANIF = os.path.join(BASE, "data", "planillas", "Planificación Financiera 2026.xlsx")
PAYROLL = os.path.join(BASE, "data", "planillas", "Sueldos_Febrero_2026.xlsx")
CHARTS_DIR = os.path.join(BASE, "data", "outputs", "charts")
OUTPUT_DOCX = os.path.join(BASE, "data", "outputs", "Analisis_Financiero_2027_UnionX_v8.docx")
os.makedirs(CHARTS_DIR, exist_ok=True)

AZUL = "#5B9BD5"; AZUL_OSC = "#2E75B6"; AZUL_LOGO = "#4A90E2"
GRIS_OSC = "#0F172A"; GRIS_MED = "#475569"; GRIS_CLR = "#CBD5E1"
VERDE = "#10B981"; VERDE_OSC = "#059669"
ROJO = "#DC2626"; NARANJA = "#EA580C"; AMARILLO = "#FACC15"
FONDO = "#F8FAFC"; BLANCO = "#FFFFFF"

plt.rcParams.update({
    'font.family': 'Arial', 'font.size': 10, 'axes.titlesize': 12, 'axes.titleweight': 'bold',
    'axes.labelsize': 10, 'axes.spines.top': False, 'axes.spines.right': False,
    'axes.edgecolor': GRIS_MED, 'xtick.color': GRIS_MED, 'ytick.color': GRIS_MED,
    'text.color': GRIS_OSC, 'figure.facecolor': 'white', 'axes.facecolor': 'white',
})

# ===== LECTURA DATOS =====
print("📥 Leyendo datos...")
wb = openpyxl.load_workbook(PLANIF, data_only=True)
ws_pl = wb['P&L']
cols_2026 = [c for c in range(7, ws_pl.max_column + 1)
             if ws_pl.cell(3, c).value and hasattr(ws_pl.cell(3, c).value, 'year') and ws_pl.cell(3, c).value.year == 2026]
cols_2025 = [c for c in range(7, ws_pl.max_column + 1)
             if ws_pl.cell(3, c).value and hasattr(ws_pl.cell(3, c).value, 'year') and ws_pl.cell(3, c).value.year == 2025]
MESES = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']

def serie_pl(fila):
    return [ws_pl.cell(fila, c).value or 0 for c in cols_2026]

ventas_m = serie_pl(4); cd_m = [-v for v in serie_pl(7)]; mf_m = serie_pl(9)
com_m = [-v for v in serie_pl(13)]; flete_m = [-v for v in serie_pl(14)]; mkt_perf_m = [-v for v in serie_pl(15)]
oce_m = [com_m[i] + flete_m[i] + mkt_perf_m[i] for i in range(12)]
mc_m = serie_pl(18); sueldos_m = [-v for v in serie_pl(22)]; oficina_m = [-v for v in serie_pl(23)]
suscripciones_m = [-v for v in serie_pl(33)]; mkt_branding_m = [-v for v in serie_pl(36)]
deprec_m = [-v for v in serie_pl(38)]; gav_m = [-v for v in serie_pl(40)]
ebit_m = serie_pl(42); gno_m = [-v for v in serie_pl(52)]; rai_m = serie_pl(55); rdi_m = serie_pl(57)

# Anuales 2025 y 2026
V25 = sum((ws_pl.cell(4, c).value or 0) for c in cols_2025)
MC25 = sum((ws_pl.cell(18, c).value or 0) for c in cols_2025)
V26 = sum(ventas_m); CD26 = sum(cd_m); MF26 = sum(mf_m)
OCE26 = sum(oce_m); MC26 = sum(mc_m)
SUELDOS26 = sum(sueldos_m); OFICINA26 = sum(oficina_m)
SUSCRIPCIONES26 = sum(suscripciones_m); MKT_BRANDING26 = sum(mkt_branding_m)
DEPREC26 = sum(deprec_m); GAV26 = sum(gav_m)
EBIT26 = sum(ebit_m); EBITDA26 = EBIT26 + DEPREC26
GNO26 = sum(gno_m); RAI26 = sum(rai_m); UN26 = sum(rdi_m)
MFp26 = MF26/V26; MCp26 = MC26/V26
MCp25 = MC25/V25 if V25 else 0
CREC_25_26 = V26/V25 - 1 if V25 else 0

# KT mensual
ws_kt = wb['KT']
cols_kt = [c for c in range(1, ws_kt.max_column + 1)
           if ws_kt.cell(7, c).value and hasattr(ws_kt.cell(7, c).value, 'year') and ws_kt.cell(7, c).value.year == 2026]
existencias_m = [ws_kt.cell(9, c).value or 0 for c in cols_kt]
cxc_m = [ws_kt.cell(10, c).value or 0 for c in cols_kt]
cxp_m = [abs(ws_kt.cell(15, c).value or 0) for c in cols_kt]
meses_exist_m = [ws_kt.cell(23, c).value or 0 for c in cols_kt]
kt_neto_m = [ws_kt.cell(19, c).value or 0 for c in cols_kt]
EXIST_PROM = sum(existencias_m)/12; EXIST_DIC = existencias_m[-1]
MESES_EXIST_PROM = sum(meses_exist_m)/12

# Deuda
ws_df = wb['Deuda financiera']
cols_df = [c for c in range(1, ws_df.max_column + 1)
           if ws_df.cell(4, c).value and hasattr(ws_df.cell(4, c).value, 'year') and ws_df.cell(4, c).value.year == 2026]
deuda_total_m = [ws_df.cell(11, c).value or 0 for c in cols_df]
credito_pesos_m = [ws_df.cell(32, c).value or 0 for c in cols_df]
comex_m = [ws_df.cell(45, c).value or 0 for c in cols_df]
DEUDA_TOTAL_DIC = deuda_total_m[-1]
CREDITO_PESOS_DIC = credito_pesos_m[-1]
COMEX_DIC = comex_m[-1]
PATRIM26 = 683924.340
TASA_INT_EFEC = GNO26 / ((1954325 + DEUDA_TOTAL_DIC) / 2)
DE26 = DEUDA_TOTAL_DIC / PATRIM26

# Payroll
df_pay = pd.read_excel(PAYROLL, sheet_name='LIBRO', header=1)
df_pay = df_pay.rename(columns={df_pay.columns[3]: 'NOMBRE', df_pay.columns[4]: 'LN',
                                df_pay.columns[5]: 'AREA', df_pay.columns[6]: 'SUBAREA',
                                df_pay.columns[7]: 'CARGO'})
df_pay['TH'] = pd.to_numeric(df_pay['TOTAL HABERES'], errors='coerce').fillna(0)
df_pay['FECHA_CTO'] = pd.to_datetime(df_pay.iloc[:, 9], errors='coerce')
df_pay_real = df_pay[df_pay['CARGO'].notna() & (df_pay['CARGO'] != '')].copy()

# Clasificación corregida: Control Gestión NO es gerencia
def clasif_nivel(cargo, nombre):
    c = str(cargo).upper(); n = str(nombre).upper()
    if 'NOVOA MATTE' in n: return '1.1 - CEO (Gerencia General)'
    if 'NEUENSCHWANDER' in n: return '1.2 - Co-Founder (vitalicio)'
    if 'BROWNE URZUA' in n: return '1.3 - Ger. Finanzas y Operaciones'
    if 'VASQUEZ ZENTENO' in n: return '1.4 - Ger. Comercial'
    if 'GUZMÁN URZÚA' in n or 'GUZMAN URZUA' in n: return '1.5 - Ger. Productos'
    if 'PASTRAN ORTEGA' in n: return '2.1 - Sub-gerencia Control de Gestión'  # Reclasificado
    if 'SUB GERENCIA' in c: return '2.2 - Sub-gerencia'
    if 'JEFATURA' in c: return '3 - Jefatura'
    if 'PLANNER' in c or 'COORDIN' in c: return '4 - Coordinación/Planner'
    if c == 'KAM': return '5 - KAM'
    if c == 'ANALISTA': return '6 - Analista'
    if c == 'ADMINISTRATIVO': return '7 - Operativo Administrativo'
    if c in ['OPERARIO', 'RUTERO']: return '8 - Operativo Bodega/Reparto'
    if 'DISE' in c: return '9 - Diseño'
    return '10 - Otro'

df_pay_real['NIVEL'] = df_pay_real.apply(lambda r: clasif_nivel(r['CARGO'], r['NOMBRE']), axis=1)

# Indemnizaciones (para escenario cambio operador agosto 2026)
fecha_corte = pd.Timestamp('2026-08-01')
df_pay_real['ANIOS_SERVICIO'] = (fecha_corte - df_pay_real['FECHA_CTO']).dt.days / 365.25
df_pay_real['ANIOS_SERVICIO'] = df_pay_real['ANIOS_SERVICIO'].clip(lower=1, upper=11)
df_pay_real['INDEMN'] = df_pay_real['TH'] * df_pay_real['ANIOS_SERVICIO']

# Operativos que salen en escenario cambio operador
jefaturas_integrar = ['BELLOLIO', 'ORTEGA MOLINA', 'GRISMAN']
m_op = df_pay_real['AREA'].astype(str).str.upper() == 'OPERACIONES'
m_no_jef = ~df_pay_real['NOMBRE'].astype(str).str.upper().str.contains('|'.join(jefaturas_integrar), na=False)
df_operativos_salen = df_pay_real[m_op & m_no_jef]
TH_OP_SALEN = df_operativos_salen['TH'].sum()
INDEMN_OP_SALEN = df_operativos_salen['INDEMN'].sum()
N_OP_SALEN = len(df_operativos_salen)
TH_JEFATURAS_INTEGRAR = df_pay_real[df_pay_real['NOMBRE'].astype(str).str.upper().str.contains('|'.join(jefaturas_integrar), na=False)]['TH'].sum()
# Gerente de Operaciones (Andrés Browne) — se mantiene también al ir a operador externo
TH_GERENTE_OPS = df_pay_real[df_pay_real['NOMBRE'].astype(str).str.upper().str.contains('BROWNE URZUA', na=False)]['TH'].sum()
# Asumimos 50% del sueldo del Gerente de Finanzas y Operaciones se asigna a Operaciones (otro 50% a Finanzas)
TH_GERENTE_OPS_ASIGNADO = TH_GERENTE_OPS * 0.50
# Total jefaturas retenidas al ir a operador
TH_RETENIDO_OPERADOR = TH_JEFATURAS_INTEGRAR + TH_GERENTE_OPS_ASIGNADO

print(f"  Venta 2025: ${V25/1000:,.0f} MM | 2026: ${V26/1000:,.0f} MM | Crec 25→26: {CREC_25_26*100:+.1f}%")
print(f"  MC% 2025: {MCp25*100:.1f}% | 2026: {MCp26*100:.1f}%")
print(f"  Operativos que salen: n={N_OP_SALEN}, TH=${TH_OP_SALEN/1e6:.1f}MM/m, Indem=${INDEMN_OP_SALEN/1e6:.0f}MM")
print(f"  Jefaturas integradas: ${TH_JEFATURAS_INTEGRAR/1e6:.2f}MM/m (Max+Gerardo+Yohana)")
print(f"  Gerente Ops (50% asignado): ${TH_GERENTE_OPS_ASIGNADO/1e6:.2f}MM/m")
print(f"  Total retenido modelo operador: ${TH_RETENIDO_OPERADOR/1e6:.2f}MM/m (${TH_RETENIDO_OPERADOR*12*1.12/1e6:.0f} MM/año c/cargas)")


# ===== PARÁMETROS MODELO 2027 =====
TASA_IMP = 0.27
TASA_INT = TASA_INT_EFEC
UF = 0.04
AJUSTE_SAL = 0.12
# Plan IA SALIDAS — SIN Gabriela (se mantiene)
SALIDAS_PLAN_IA_5 = {
    'Analista Log. Inversa': 1004852,
    'Analista Contable #1': 1015754,
    'Analista Contable #2': 770000,
    'Facturadora #1': 895815,
    'Facturadora #2': 822212,
}
# Camila Villalta SE MANTIENE (no sale en plan IA 2027). Se incluye como posible eficiencia futura no comprometida.
SAL_CAMILA_FUTURO = 1834511  # Solo referencia, no se aplica a base 2027
total_salidas_anual = sum(SALIDAS_PLAN_IA_5.values()) * 12 / 1000

# Costo operativo REAL (de imagen KPI app)
COSTO_X_PED_ACTUAL = 2812
PEDIDOS_ANO = 180000
AOV = 36904
COSTO_OP_ANUAL_26 = COSTO_X_PED_ACTUAL * PEDIDOS_ANO / 1000  # = $506 MM

# Cargas patronales
CARGAS = 1.12

# Reducción mensual por salidas plan IA OPERATIVAS (afectan costo/pedido)
# Solo 3 personas afectan operación: Jorgelis (Post Venta), Iris (Facturación), Stipp (Facturación)
red_men_op_jul_full = 1004852  # Jorgelis sale 31-jul
red_men_op_ago_med = 895815 / 2  # Iris sale 15-ago (medio mes)
red_men_op_ago_full = 895815  # sept en adelante
red_men_op_sep_med = 822212 / 2  # Stipp sale 15-sep
red_men_op_sep_full = 822212  # oct en adelante

# Costo mensual in-house 2026 con plan IA
costo_men_inh = []
costo_base_men = COSTO_OP_ANUAL_26 * 1000 / 12  # en pesos
for m in range(1, 13):
    cm = costo_base_men
    if m >= 8: cm -= red_men_op_jul_full * CARGAS  # Jorgelis
    if m == 8: cm += 0  # ago Iris aún
    if m == 9: cm -= red_men_op_ago_med * CARGAS  # set Iris medio
    if m >= 10: cm -= red_men_op_ago_full * CARGAS
    if m == 10: cm -= red_men_op_sep_med * CARGAS
    if m >= 11: cm -= red_men_op_sep_full * CARGAS
    costo_men_inh.append(cm)
COSTO_INH_2026_TOTAL = sum(costo_men_inh) / 1000  # miles
costo_x_ped_proy_men = [c / (PEDIDOS_ANO/12) for c in costo_men_inh]
COSTO_X_PED_POST_SALIDAS = costo_x_ped_proy_men[-1]  # diciembre = full effect

# Benchmarks operador
BENCH_OP_OPT = 1500
BENCH_OP_BASE = 2000
BENCH_OP_CONS = 2500
COSTO_JEFATURAS_INTEG = TH_RETENIDO_OPERADOR * 12 * CARGAS / 1000  # miles anual (incluye Gerente Ops asignado 50%)

# ===== FUNCIONES 2027 =====
def calc_suscripciones_2027():
    yuju = 7360; multivende = 0; odoo = 4100
    claude = 2250 * 12 * 0.920; chatgpt = 980 * 1.04
    otros = max(5000, (SUSCRIPCIONES26 - 7360 - 1980 - 9570 - 2500 - 980) * 1.04)
    return yuju + odoo + claude + chatgpt + otros

# Sueldos 2027 — Gabriela SE MANTIENE (no en salidas)
def calc_gav_2027(venta, ef_gav=0.0):
    sueldos_27 = (SUELDOS26 - total_salidas_anual) * (1 + AJUSTE_SAL)
    # NOTA: total_salidas_anual ya NO incluye a Gabriela (no está en SALIDAS_PLAN_IA_5)
    oficina_27 = OFICINA26 * (1 + UF)
    susc_27 = calc_suscripciones_2027()
    mkt_b_27 = 0.03 * venta
    deprec_27 = DEPREC26
    otros_27 = (GAV26 - SUELDOS26 - OFICINA26 - SUSCRIPCIONES26 - MKT_BRANDING26 - DEPREC26) * (1 + UF)
    total = (sueldos_27 + oficina_27 + susc_27 + mkt_b_27 + deprec_27 + otros_27) * (1 - ef_gav)
    return {'sueldos': sueldos_27, 'oficina': oficina_27, 'suscripciones': susc_27,
            'mkt_branding': mkt_b_27, 'depreciacion': deprec_27, 'otros': otros_27, 'total': total}

def calc_pyl_2027(venta, mc_pct=MCp26, ef_gav=0.0, deuda=DEUDA_TOTAL_DIC):
    mc = mc_pct * venta
    gav = calc_gav_2027(venta, ef_gav)
    ebit = mc - gav['total']
    ebitda = ebit + gav['depreciacion']
    gno = TASA_INT * deuda
    rai = ebit - gno
    impuesto = max(0, rai * TASA_IMP)
    un = rai - impuesto
    return {'venta': venta, 'mc': mc, 'mc_pct': mc_pct, 'gav': gav, 'gav_total': gav['total'],
            'ebit': ebit, 'ebit_pct': ebit/venta if venta else 0,
            'ebitda': ebitda, 'ebitda_pct': ebitda/venta if venta else 0,
            'gno': gno, 'rai': rai, 'impuesto': impuesto, 'un': un, 'deuda': deuda}

def venta_para_ebitda(target_pct, mc_pct=MCp26, ef_gav=0.0):
    gav0 = calc_gav_2027(0, ef_gav)
    deprec = gav0['depreciacion']
    coef = mc_pct - 0.03 - target_pct
    if coef <= 0: return None
    return (gav0['total'] - deprec) / coef

def venta_para_un(target_un, mc_pct=MCp26, ef_gav=0.0):
    rai = target_un / (1 - TASA_IMP)
    ebit = rai + TASA_INT * DEUDA_TOTAL_DIC
    gav0 = calc_gav_2027(0, ef_gav)
    return (ebit + gav0['total']) / (mc_pct - 0.03)

# Base 2027 sin cambios
V_BASE_EBITDA_27 = venta_para_ebitda(0.10, mc_pct=MCp26, ef_gav=0)

# ===== MATRIZ EXTENDIDA DE ESCENARIOS =====
# 7 niveles de MC (27% → 33%), con ef GAV gradual y política inv gradual
matriz_esc = []
for i, mc in enumerate([0.27, 0.28, 0.29, 0.30, 0.31, 0.32, 0.33]):
    # Asumir ef GAV gradual y política inv gradual (más palancas = más esfuerzo)
    ef_gav = i * 0.025  # 0%, 2.5%, 5%, 7.5%, 10%, 12.5%, 15%
    pol_inv = 4.4 - (i * 0.15)  # 4.4 → 3.5 m
    v_req = venta_para_ebitda(0.10, mc_pct=mc, ef_gav=ef_gav)
    if v_req:
        crec = (v_req/V26 - 1) * 100
        pyl = calc_pyl_2027(v_req, mc_pct=mc, ef_gav=ef_gav)
        un = pyl['un']
        # Sensibilidad pedidos
        pedidos = PEDIDOS_ANO * (v_req / V26)
    else:
        crec = None; pyl = None; un = None; pedidos = None
    matriz_esc.append({
        'mc': mc, 'ef_gav': ef_gav, 'pol_inv': pol_inv,
        'venta': v_req, 'crec': crec, 'un': un, 'pyl': pyl, 'pedidos': pedidos
    })

# 5 escenarios "Top" para tabla comparativa principal
# Base = continuidad histórica
esc_top = [
    {'nombre': 'Status quo (sin cambios)', 'mc': MCp26, 'ef_gav': 0.00, 'crec': 0.00, 'pol_inv': 4.4},
    {'nombre': 'Realista (continuidad histórica)', 'mc': 0.29, 'ef_gav': 0.05, 'crec': CREC_25_26, 'pol_inv': 4.0},
    {'nombre': 'Mejorado (plan IA + 5 palancas)', 'mc': 0.31, 'ef_gav': 0.08, 'crec': 0.12, 'pol_inv': 4.0},
    {'nombre': 'Optimista (palancas plenas)', 'mc': 0.33, 'ef_gav': 0.12, 'crec': 0.17, 'pol_inv': 3.5},
    {'nombre': 'Ambicioso (stretch)', 'mc': 0.35, 'ef_gav': 0.15, 'crec': 0.25, 'pol_inv': 3.5},
]
for e in esc_top:
    v = V26 * (1 + e['crec'])
    # Efecto política inventario sobre deuda: si pol_inv cambia respecto a 2026 (4,4 m promedio), modifica COMEX
    costo_venta_e = v * (CD26/V26)
    inventario_pol = costo_venta_e * e['pol_inv'] / 12 * 1.10  # incluye 10% obsoleto
    delta_inventario = inventario_pol - EXIST_PROM  # diferencia vs 2026
    deuda_27 = max(800000, DEUDA_TOTAL_DIC + delta_inventario)  # piso $800 MM (no todo es COMEX)
    e['deuda_27'] = deuda_27
    e['delta_inventario_kt'] = delta_inventario
    pyl = calc_pyl_2027(v, mc_pct=e['mc'], ef_gav=e['ef_gav'], deuda=deuda_27)
    e['venta'] = v
    e['pedidos'] = PEDIDOS_ANO * (v / V26)
    e['ebitda'] = pyl['ebitda']
    e['ebitda_pct'] = pyl['ebitda_pct']
    e['un'] = pyl['un']
    e['gno'] = pyl['gno']
    e['retiros'] = pyl['un'] * 0.30
    e['cumple_ebitda'] = pyl['ebitda_pct'] >= 0.10
    e['cumple_un'] = pyl['un'] >= 333000
    patrim = PATRIM26 + pyl['un'] - min(pyl['un']*0.30, 100000)
    e['patrim_fin'] = patrim
    e['de'] = deuda_27 / patrim if patrim > 0 else 999
    e['cumple_de'] = 2.5 <= e['de'] <= 3.5
    # Costo/pedido in-house 2027 (asume escala con eficiencia operacional)
    # Costo total operativo escalado por ped + eficiencia
    costo_op_27 = (COSTO_OP_ANUAL_26 - total_salidas_anual * 0.4) * (1 + UF) * (1 - e['ef_gav'])
    # Pedidos escala con venta
    costo_op_27_escala = costo_op_27 * (1 + 0.5 * (v/V26 - 1))  # 50% del crec es marginal
    e['costo_inhouse_27'] = costo_op_27_escala
    e['cxp_inhouse_27'] = costo_op_27_escala / e['pedidos'] * 1000
    # Costo operador (todo-incluido + jefaturas integradas)
    costo_op_27_base = COSTO_JEFATURAS_INTEG * (1 + UF) + BENCH_OP_BASE * e['pedidos'] / 1000
    costo_op_27_opt = COSTO_JEFATURAS_INTEG * (1 + UF) + BENCH_OP_OPT * e['pedidos'] / 1000
    e['costo_op_27_base'] = costo_op_27_base
    e['cxp_op_27_base'] = costo_op_27_base / e['pedidos'] * 1000
    e['costo_op_27_opt'] = costo_op_27_opt
    e['cxp_op_27_opt'] = costo_op_27_opt / e['pedidos'] * 1000

# ===== ESCENARIO CAMBIO OPERADOR 2026 (CORREGIDO) =====
# In-house TOTAL 2026 con plan IA aplicado: $COSTO_INH_2026_TOTAL
# Cambio agosto: 5 meses operador (75K pedidos ago-dic)
PEDIDOS_AGO_DIC_26 = PEDIDOS_ANO * 5/12  # 75K
# Costo ene-jul in-house: 7 meses con costo proyectado mensual
costo_in_ene_jul = sum(costo_men_inh[:7]) / 1000  # miles
# Costo ago-dic in-house: 5 meses con costo proyectado mensual (con salidas plan IA)
costo_in_ago_dic = sum(costo_men_inh[7:]) / 1000

def costo_op_ago_dic_2026(bench):
    # Tarifa todo-incluido × pedidos
    var = bench * PEDIDOS_AGO_DIC_26 / 1000
    # Jefaturas integradas pagadas por nosotros (5 meses) — incluye Gerente Ops asignado
    fija_jef = TH_RETENIDO_OPERADOR * 5 * CARGAS / 1000
    return var + fija_jef

cambio_2026 = []
for bench, label in [(BENCH_OP_OPT, 'Optimista'), (BENCH_OP_BASE, 'Base'), (BENCH_OP_CONS, 'Conservador')]:
    costo_op = costo_op_ago_dic_2026(bench)
    total_con_cambio = costo_in_ene_jul + costo_op + INDEMN_OP_SALEN/1000
    total_sin_cambio = costo_in_ene_jul + costo_in_ago_dic
    delta = total_sin_cambio - total_con_cambio
    cambio_2026.append({
        'bench': bench, 'label': label,
        'costo_in_ene_jul': costo_in_ene_jul,
        'costo_op_ago_dic': costo_op,
        'indem': INDEMN_OP_SALEN/1000,
        'total_con_cambio': total_con_cambio,
        'total_sin_cambio': total_sin_cambio,
        'delta_neto': delta,
        'cxp_efectivo': (costo_op_ago_dic_2026(bench) * 1000) / PEDIDOS_AGO_DIC_26,  # CLP/ped
    })

# ===== PUNTO DE EQUILIBRIO LARGO PLAZO =====
# Modelo: Costo in-house = F_in + V_in × P
#         Costo operador = J_op + B_op × P
# P_eq = (J_op - F_in) / (V_in - B_op)
# Asumo 45% fijo / 55% variable del costo operativo
PCT_FIJO_INH = 0.45
F_in_actual = COSTO_OP_ANUAL_26 * PCT_FIJO_INH * 1000  # pesos
V_in_actual = COSTO_OP_ANUAL_26 * (1-PCT_FIJO_INH) / PEDIDOS_ANO * 1000  # CLP/ped
F_in_plan_ia = (COSTO_OP_ANUAL_26 - total_salidas_anual * 0.4) * PCT_FIJO_INH * 1000
V_in_plan_ia = (COSTO_OP_ANUAL_26 - total_salidas_anual * 0.4) * (1-PCT_FIJO_INH) / PEDIDOS_ANO * 1000
F_in_eficiente = F_in_plan_ia * 0.85  # eficiencia operativa adicional 15% en fijos
V_in_eficiente = V_in_plan_ia * 0.95  # -5% variable

J_op = COSTO_JEFATURAS_INTEG * 1000  # pesos anuales

def punto_eq(F_in, V_in, J_op, B_op):
    if V_in == B_op: return None
    return (J_op - F_in) / (V_in - B_op)

casos_eq = [
    {'nombre': 'Caso A — Costo in-house actual ($2.812/ped)', 'F': F_in_actual, 'V': V_in_actual},
    {'nombre': 'Caso B — Con plan IA aplicado (~$2.610/ped)', 'F': F_in_plan_ia, 'V': V_in_plan_ia},
    {'nombre': 'Caso C — Plan IA + eficiencia operacional (~$2.300/ped)', 'F': F_in_eficiente, 'V': V_in_eficiente},
]
for caso in casos_eq:
    caso['eq_opt'] = punto_eq(caso['F'], caso['V'], J_op, BENCH_OP_OPT)
    caso['eq_base'] = punto_eq(caso['F'], caso['V'], J_op, BENCH_OP_BASE)
    caso['eq_cons'] = punto_eq(caso['F'], caso['V'], J_op, BENCH_OP_CONS)

# ===== EFECTO OPERADOR LOGÍSTICO POR ESCENARIO (NUEVO v4) =====
# Para cada escenario calcular: P&L bajo in-house default vs cambio operador (base y opt)
# El cambio operador modifica el costo operativo en el GAV (sueldos OP + arriendo + insumos)
# Costo operativo in-house 2027 (full plan IA): COSTO_OP_2027_INH
COSTO_OP_2027_INH_BASE = (COSTO_OP_ANUAL_26 - total_salidas_anual * 0.4) * (1 + UF)
# = $489 MM aprox

# Componentes del costo operativo in-house que se reemplazan al ir con operador
# Mantenemos: jefaturas (Max+Gerardo+Yohana) integradas + 50% arriendo bodega (inventario)
ARRIENDO_BODEGA_ANUAL = 100000  # $100 MM/año (parte del oficina/arriendos asignable)
ARRIENDO_RETENIDO_PCT = 0.50
SUSCRIPCIONES_OP_RETENIDAS = 5000  # $5 MM/año mantenido
COSTO_RETENIDO_INH = (TH_JEFATURAS_INTEGRAR * 12 * CARGAS / 1000) + (ARRIENDO_BODEGA_ANUAL * ARRIENDO_RETENIDO_PCT) + SUSCRIPCIONES_OP_RETENIDAS

def costo_op_27_escala(venta, ef_gav):
    """Costo operativo in-house 2027 escalado con venta. 50% del crec marginal escala."""
    factor = 1 + 0.5 * (venta/V26 - 1)
    return COSTO_OP_2027_INH_BASE * factor * (1 - ef_gav)

def costo_operador_27(pedidos, bench):
    """Costo modelo operador externo: tarifa todo-incluida + jefaturas retenidas + arriendo bodega 50% para inventario."""
    return bench * pedidos / 1000 + COSTO_RETENIDO_INH * (1 + UF)

# Para cada escenario top, calcular bajo 3 estructuras: in-house, operador base, operador opt
for e in esc_top:
    v = e['venta']
    pedidos_e = e['pedidos']
    ef_gav = e['ef_gav']
    # Costo operativo in-house
    c_inh = costo_op_27_escala(v, ef_gav)
    # Costo operador (base y opt)
    c_op_b = costo_operador_27(pedidos_e, BENCH_OP_BASE)
    c_op_o = costo_operador_27(pedidos_e, BENCH_OP_OPT)
    # Delta operativo por estructura (operador vs in-house)
    e['delta_op_base'] = c_inh - c_op_b  # positivo = operador ahorra
    e['delta_op_opt'] = c_inh - c_op_o
    # Recalcular EBITDA y UN bajo cada estructura
    # delta_op_base y delta_op_opt están en miles (mismo que ebitda_inh) — NO convertir
    ebitda_inh = e['ebitda']
    ebitda_op_base = ebitda_inh + e['delta_op_base']
    ebitda_op_opt = ebitda_inh + e['delta_op_opt']
    e['ebitda_op_base'] = ebitda_op_base
    e['ebitda_op_opt'] = ebitda_op_opt
    e['ebitda_pct_op_base'] = ebitda_op_base / v
    e['ebitda_pct_op_opt'] = ebitda_op_opt / v
    # UN bajo cada estructura
    ebit_op_base = ebitda_op_base - DEPREC26
    ebit_op_opt = ebitda_op_opt - DEPREC26
    rai_op_base = ebit_op_base - TASA_INT * DEUDA_TOTAL_DIC
    rai_op_opt = ebit_op_opt - TASA_INT * DEUDA_TOTAL_DIC
    e['un_op_base'] = rai_op_base * (1 - TASA_IMP) if rai_op_base > 0 else rai_op_base
    e['un_op_opt'] = rai_op_opt * (1 - TASA_IMP) if rai_op_opt > 0 else rai_op_opt
    # Cumplimientos bajo cada estructura
    e['cumple_ebitda_op_base'] = ebitda_op_base / v >= 0.10
    e['cumple_ebitda_op_opt'] = ebitda_op_opt / v >= 0.10
    e['cumple_un_op_base'] = e['un_op_base'] >= 333000
    e['cumple_un_op_opt'] = e['un_op_opt'] >= 333000
    # Costo operativo total para tabla
    e['costo_op_inh'] = c_inh
    e['costo_op_base_27'] = c_op_b
    e['costo_op_opt_27'] = c_op_o

# ===== SENSIBILIDADES =====
sens_2026 = []
for delta in [-0.20, -0.10, 0, 0.10, 0.20]:
    v_alt = V26 * (1 + delta)
    mc_alt = MCp26 * v_alt
    gav_fijo = GAV26 - MKT_BRANDING26
    mkt_b_alt = (MKT_BRANDING26/V26) * v_alt
    gav_alt = gav_fijo + mkt_b_alt
    ebit_alt = mc_alt - gav_alt
    ebitda_alt = ebit_alt + DEPREC26
    rai_alt = ebit_alt - GNO26
    un_alt = rai_alt * (1 - TASA_IMP) if rai_alt > 0 else rai_alt
    patrim_27 = PATRIM26 + un_alt - 100000
    sens_2026.append({'delta': delta, 'venta26': v_alt, 'un26': un_alt,
                       'patrim_27': patrim_27,
                       'de_27': DEUDA_TOTAL_DIC/patrim_27 if patrim_27>0 else 999})

# Plan contingencia
gav27_base = calc_pyl_2027(V_BASE_EBITDA_27)['gav_total']
capas_cont = [
    {'capa': 1, 'nombre': 'Discrecionales (más lejano)', 'color': VERDE,
     'items': ['Marketing Branding (3% V)', 'Asesorías opcionales', 'I+D no comprometido',
               'Capacitación', 'Viajes', 'Beneficios extra-legales'], 'pct': 0.10, 'vel': 'Inmediata'},
    {'capa': 2, 'nombre': 'SaaS optimizable', 'color': VERDE,
     'items': ['Yuju (50%)', 'Suscripciones secundarias', 'Software duplicado'],
     'pct': 0.03, 'vel': '1-2 meses'},
    {'capa': 3, 'nombre': 'Admin soporte', 'color': AMARILLO,
     'items': ['Plan IA acelerado', 'Asesorías externas', 'Mandos medios admin'],
     'pct': 0.05, 'vel': '2-3 meses'},
    {'capa': 4, 'nombre': 'Operativos no críticos', 'color': NARANJA,
     'items': ['Tercerización marginal peak', 'Mkt Performance -30%',
               'Reducción horas extra', 'Pausar PP&E'],
     'pct': 0.05, 'vel': '3-4 meses'},
    {'capa': 5, 'nombre': 'Estructura operativa y mandos medios', 'color': ROJO,
     'items': ['Cambio a operador logístico', 'Jefaturas operativas (Logística, Post Venta, Facturación)',
               'Sub-gerencia Control Gestión', 'Renegociación arriendos'],
     'pct': 0.08, 'vel': '4-6 meses'},
    {'capa': 6, 'nombre': 'CORE inviolable (mínimo)', 'color': GRIS_OSC,
     'items': ['CEO (Martín Novoa)', 'Co-Founder Erich (vitalicio)',
               'Gerencias Comercial · Productos · Finanzas',
               'KAMs comerciales activos'],
     'pct': 0.00, 'vel': 'NO TOCAR'},
]
for capa in capas_cont:
    capa['monto'] = gav27_base * capa['pct']


# =====================================================================
# GRÁFICOS
# =====================================================================
print("\n📈 Generando gráficos...")

def save(name):
    path = os.path.join(CHARTS_DIR, name + '.png')
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return path

# G1: Venta mensual
fig, ax = plt.subplots(figsize=(10, 4.5))
x = range(12)
ax.bar(x, [v/1000 for v in ventas_m], color=AZUL, label='Venta', alpha=0.85, width=0.7)
ax.plot(x, [v/1000 for v in mc_m], color=VERDE, marker='o', linewidth=2.5, label='MC')
ax.plot(x, [v/1000 for v in ebit_m], color=NARANJA, marker='s', linewidth=2, label='EBIT')
ax.axhline(0, color=GRIS_MED, linewidth=0.5)
ax.set_xticks(x); ax.set_xticklabels(MESES); ax.set_ylabel('MM CLP')
ax.set_title('Evolución mensual 2026 — Venta, MC y EBIT', loc='left')
ax.legend(loc='upper left', frameon=False); ax.grid(True, axis='y', alpha=0.3)
save('01_venta_mensual')

# G2: GAV donut
fig, ax = plt.subplots(figsize=(7, 5))
comp = [('Sueldos', SUELDOS26, AZUL_OSC), ('Oficina/Arriendos', OFICINA26, AZUL),
        ('Suscripciones', SUSCRIPCIONES26, AZUL_LOGO), ('Mkt Branding', MKT_BRANDING26, NARANJA),
        ('Depreciación', DEPREC26, GRIS_MED),
        ('Otros', GAV26-SUELDOS26-OFICINA26-SUSCRIPCIONES26-MKT_BRANDING26-DEPREC26, GRIS_CLR)]
vals = [c[1]/1000 for c in comp]
labels = [f"{c[0]}\n${c[1]/1000:,.0f} MM ({c[1]/GAV26*100:.0f}%)" for c in comp]
wedges, _ = ax.pie(vals, colors=[c[2] for c in comp], startangle=90, wedgeprops=dict(width=0.4))
ax.legend(wedges, labels, loc='center left', bbox_to_anchor=(1.05, 0.5), frameon=False, fontsize=9)
ax.set_title(f'Estructura GAV 2026 — Total ${GAV26/1000:,.0f} MM', loc='left')
save('02_gav_estructura')

# G3: Waterfall
fig, ax = plt.subplots(figsize=(11, 5))
pasos = [('Venta', V26/1000, AZUL), ('- Costo Directo', -CD26/1000, ROJO),
         ('Margen Frontal', MF26/1000, AZUL_OSC), ('- Otros Costos', -OCE26/1000, ROJO),
         ('Margen Contrib.', MC26/1000, AZUL_OSC), ('- GAV', -GAV26/1000, ROJO),
         ('EBIT', EBIT26/1000, VERDE), ('- GNO', -GNO26/1000, ROJO),
         ('UN', UN26/1000, VERDE_OSC)]
cumul = V26/1000
for i, (nom, val, color) in enumerate(pasos):
    if nom in ('Margen Frontal', 'Margen Contrib.', 'EBIT', 'UN', 'Venta'):
        ax.bar(i, val, bottom=0, color=color, width=0.6)
        ax.text(i, val + 100, f'${val:,.0f}', ha='center', va='bottom', fontweight='bold', fontsize=9)
        if nom != 'Venta': cumul = val
    else:
        ax.bar(i, val, bottom=cumul if val < 0 else cumul - val, color=color, width=0.6)
        cumul = cumul + val
        ax.text(i, cumul + (val/2 if val<0 else -val/2), f'${val:,.0f}', ha='center', va='center', color='white', fontweight='bold', fontsize=9)
ax.set_xticks(range(len(pasos))); ax.set_xticklabels([p[0] for p in pasos], rotation=35, ha='right')
ax.set_ylabel('MM CLP'); ax.set_title('Cascada P&L 2026', loc='left')
ax.grid(True, axis='y', alpha=0.3); ax.spines['bottom'].set_visible(False)
save('03_waterfall')

# G4: Matriz extendida MC% → Venta requerida (NUEVO v3)
fig, ax = plt.subplots(figsize=(11, 5.5))
mcs = [e['mc']*100 for e in matriz_esc]
crecs = [e['crec'] for e in matriz_esc if e['crec'] is not None]
ventas_req = [e['venta']/1000 for e in matriz_esc if e['venta'] is not None]
ef_gavs = [e['ef_gav']*100 for e in matriz_esc]
ax.plot(mcs, crecs, marker='o', linewidth=3, color=AZUL_OSC, markersize=10)
for mc, cr, v, ef in zip(mcs, crecs, ventas_req, ef_gavs):
    color_pt = VERDE if cr <= CREC_25_26*100 else NARANJA if cr <= 15 else ROJO
    ax.scatter([mc], [cr], color=color_pt, s=100, zorder=5)
    ax.annotate(f'+{cr:.0f}%\n(${v:,.0f}MM)\nef GAV {ef:.0f}%',
                xy=(mc, cr), xytext=(0, 15), textcoords='offset points',
                ha='center', fontsize=8, fontweight='bold' if mc in [29.0, 31.0, 33.0] else 'normal')
ax.axhline(CREC_25_26*100, color=VERDE, linestyle='--', alpha=0.7, label=f'Crec histórico 25→26: +{CREC_25_26*100:.1f}%')
ax.axhline(15, color=NARANJA, linestyle=':', alpha=0.5, label='Crec moderado: +15%')
ax.set_xlabel('Margen Contribución %')
ax.set_ylabel('Crecimiento venta requerido %')
ax.set_title('Matriz MC% × Crecimiento V requerido para EBITDA 10% (con ef GAV gradual)', loc='left')
ax.legend(loc='upper right', frameon=False)
ax.grid(True, alpha=0.3)
save('04_matriz_mc')

# G5: 5 escenarios top - EBITDA
fig, ax = plt.subplots(figsize=(11, 6))
nombres = [e['nombre'] for e in esc_top]
ebitdas = [e['ebitda_pct']*100 for e in esc_top]
colors_e = [ROJO if not e['cumple_ebitda'] else VERDE for e in esc_top]
bars = ax.barh(nombres, ebitdas, color=colors_e, alpha=0.85)
ax.axvline(10, color=AZUL_OSC, linestyle='--', linewidth=2, label='Objetivo CEO: 10%')
for b, esc in zip(bars, esc_top):
    ax.text(b.get_width()+0.3, b.get_y()+b.get_height()/2,
            f'EBITDA {esc["ebitda_pct"]*100:.1f}% · UN ${esc["un"]/1000:,.0f}M · Crec +{esc["crec"]*100:.0f}%',
            va='center', fontsize=9, fontweight='bold')
ax.set_xlabel('EBITDA %')
ax.set_title('5 Escenarios Top — Cumplimiento EBITDA 10%', loc='left')
ax.legend(loc='lower right', frameon=False); ax.grid(True, axis='x', alpha=0.3)
save('05_escenarios_top')

# G6: Cumplimiento matriz
fig, ax = plt.subplots(figsize=(11, 5))
restr = ['EBITDA ≥10%', 'UN ≥ $333MM', 'D/E 2,5-3,5']
nom_e = [e['nombre'] for e in esc_top]
mat = np.array([[1 if e['cumple_ebitda'] else 0,
                 1 if e['cumple_un'] else 0,
                 1 if e['cumple_de'] else 0] for e in esc_top])
import matplotlib.colors as mcolors
cmap = mcolors.ListedColormap([ROJO, VERDE])
ax.imshow(mat, aspect='auto', cmap=cmap, vmin=0, vmax=1)
ax.set_xticks(range(3)); ax.set_xticklabels(restr)
ax.set_yticks(range(len(nom_e))); ax.set_yticklabels(nom_e)
for i in range(len(nom_e)):
    for j in range(3):
        ax.text(j, i, '✓' if mat[i,j]==1 else '✗', ha='center', va='center', color='white', fontsize=22, fontweight='bold')
ax.set_title('Matriz cumplimiento — Restricciones CEO 2027', loc='left')
save('06_cumplimiento')

# G7: Sensibilidad ±20%
fig, axes = plt.subplots(1, 2, figsize=(13, 5))
deltas = [s['delta']*100 for s in sens_2026]
vts = [s['venta26']/1000 for s in sens_2026]; uns = [s['un26']/1000 for s in sens_2026]
ax = axes[0]
ax.bar(deltas, vts, color=[ROJO if d<0 else VERDE if d>0 else AZUL for d in deltas], alpha=0.85, width=8)
for d, v in zip(deltas, vts): ax.text(d, v+100, f'${v:,.0f}M', ha='center', fontweight='bold', fontsize=9)
ax.set_xlabel('Δ venta 2026 (%)'); ax.set_ylabel('Venta 2026 (MM CLP)')
ax.set_title('Venta 2026 sensibilidad', loc='left'); ax.grid(True, axis='y', alpha=0.3)
ax = axes[1]
ax.bar(deltas, uns, color=[ROJO if u<UN26/1000 else VERDE for u in uns], alpha=0.85, width=8)
for d, u in zip(deltas, uns):
    ax.text(d, u+5 if u>0 else u-15, f'${u:,.0f}M', ha='center', fontweight='bold', fontsize=9,
            color=GRIS_OSC if u>0 else ROJO)
ax.axhline(0, color=GRIS_MED, linewidth=0.5)
ax.set_xlabel('Δ venta 2026 (%)'); ax.set_ylabel('UN 2026 (MM CLP)')
ax.set_title('Apalancamiento operativo', loc='left'); ax.grid(True, axis='y', alpha=0.3)
save('07_sens_2026')

# G8: Plan contingencia
fig, ax = plt.subplots(figsize=(11, 6))
montos = [c['monto']/1000 for c in capas_cont]
colors_c = [c['color'] for c in capas_cont]
y = list(range(len(capas_cont)))
bars = ax.barh(y, montos, color=colors_c, alpha=0.85)
for b, capa in zip(bars, capas_cont):
    ax.text(b.get_width()+1, b.get_y()+b.get_height()/2, f'${capa["monto"]/1000:,.0f} MM · {capa["vel"]}',
            va='center', fontsize=9, fontweight='bold')
ax.set_yticks(y); ax.set_yticklabels([f'Capa {c["capa"]}\n{c["nombre"][:32]}' for c in capas_cont])
ax.invert_yaxis(); ax.set_xlabel('Capacidad recorte anual (MM CLP)')
ax.set_title('Plan de Contingencia — Capas desde lejano al CORE', loc='left')
ax.grid(True, axis='x', alpha=0.3)
save('08_contingencia')

# G9: Estacionalidad
fig, ax = plt.subplots(figsize=(10, 4.5))
pct = [v/V26*100 for v in ventas_m]; prom = 100/12
colors_e = [VERDE if p > prom*1.05 else ROJO if p < prom*0.95 else AZUL for p in pct]
ax.bar(MESES, pct, color=colors_e, alpha=0.85)
ax.axhline(prom, color=GRIS_OSC, linestyle='--', label=f'Promedio mensual ({prom:.1f}%)')
for m, p in zip(MESES, pct): ax.text(m, p+0.15, f'{p:.1f}%', ha='center', fontsize=8, fontweight='bold')
ax.set_ylabel('% del año'); ax.set_title('Estacionalidad 2026', loc='left')
ax.legend(frameon=False); ax.grid(True, axis='y', alpha=0.3)
save('09_estacionalidad')

# G10: Existencias mensual
fig, ax1 = plt.subplots(figsize=(11, 5))
ax2 = ax1.twinx()
exist_mm = [e/1000 for e in existencias_m]
ax1.bar(MESES, exist_mm, color=AZUL_LOGO, alpha=0.85, label='Existencias')
ax2.plot(MESES, meses_exist_m, color=NARANJA, marker='o', linewidth=2.5, label='Meses inventario')
ax2.axhline(5.5, color=ROJO, linestyle='--', alpha=0.7, label='Política 5,5 m')
ax2.axhline(4.0, color=VERDE, linestyle=':', alpha=0.7, label='Recomendación 4 m')
ax1.set_ylabel('MM CLP', color=AZUL_OSC); ax2.set_ylabel('Meses', color=NARANJA)
ax1.set_title('Existencias mensual 2026 — saldos planilla vs política', loc='left')
fig.legend(loc='upper right', bbox_to_anchor=(0.92, 0.9), frameon=False, fontsize=8)
ax1.grid(True, axis='y', alpha=0.3)
save('10_existencias')

# G11: Deuda mensual
fig, ax = plt.subplots(figsize=(11, 5))
xx = range(12)
ax.bar(xx, [c/1000 for c in credito_pesos_m], color=NARANJA, alpha=0.85, label='Crédito pesos')
ax.bar(xx, [c/1000 for c in comex_m], bottom=[c/1000 for c in credito_pesos_m], color=AZUL_OSC, alpha=0.85, label='COMEX')
for i in xx:
    tot = (credito_pesos_m[i] + comex_m[i])/1000
    ax.text(i, tot+30, f'${tot:,.0f}', ha='center', fontsize=8, fontweight='bold')
ax.set_xticks(xx); ax.set_xticklabels(MESES); ax.set_ylabel('Deuda MM CLP')
ax.set_title('Estructura de Deuda 2026 — mes a mes', loc='left')
ax.legend(frameon=False); ax.grid(True, axis='y', alpha=0.3)
save('11_deuda')

# G12: Margen mensual
fig, ax1 = plt.subplots(figsize=(11, 5))
ax2 = ax1.twinx()
mfp_m = [mf_m[i]/ventas_m[i]*100 if ventas_m[i] else 0 for i in range(12)]
mcp_m = [mc_m[i]/ventas_m[i]*100 if ventas_m[i] else 0 for i in range(12)]
com_pct = [com_m[i]/ventas_m[i]*100 if ventas_m[i] else 0 for i in range(12)]
ax1.bar([i-0.2 for i in range(12)], mfp_m, width=0.4, color=VERDE, alpha=0.85, label='MF %')
ax1.bar([i+0.2 for i in range(12)], mcp_m, width=0.4, color=AZUL_OSC, label='MC %')
ax1.set_xticks(range(12)); ax1.set_xticklabels(MESES); ax1.set_ylabel('Margen %'); ax1.set_ylim(0, 65)
ax2.plot(range(12), com_pct, marker='o', color=ROJO, linewidth=2.5, label='Comisión %')
ax2.set_ylabel('Comisión %', color=ROJO); ax2.tick_params(axis='y', labelcolor=ROJO)
ax1.set_title('Estacionalidad del margen 2026', loc='left')
fig.legend(loc='upper right', bbox_to_anchor=(0.92, 0.95), frameon=False, fontsize=8)
ax1.grid(True, axis='y', alpha=0.3)
save('12_margen_mensual')

# G13: Costo por pedido mensual (NUEVO v3 — usa $2.812)
fig, ax = plt.subplots(figsize=(11, 5))
xx = range(12)
ax.plot(xx, [c for c in costo_x_ped_proy_men], marker='o', color=AZUL_OSC, linewidth=2.5, label='In-house proyectado con plan IA')
ax.axhline(COSTO_X_PED_ACTUAL, color=GRIS_MED, linestyle=':', linewidth=1.5, label=f'In-house promedio histórico ${COSTO_X_PED_ACTUAL:,}')
ax.axhline(BENCH_OP_OPT + COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO, color=VERDE, linestyle='--', linewidth=1.5,
           label=f'Operador OPT total ${BENCH_OP_OPT + int(COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO):,} (tarifa $1.500 + jef.)')
ax.axhline(BENCH_OP_BASE + COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO, color=NARANJA, linestyle='--', linewidth=1.5,
           label=f'Operador BASE total ${BENCH_OP_BASE + int(COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO):,} (tarifa $2.000 + jef.)')
ax.axhline(BENCH_OP_CONS + COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO, color=ROJO, linestyle='--', linewidth=1.5,
           label=f'Operador CONS total ${BENCH_OP_CONS + int(COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO):,} (tarifa $2.500 + jef.)')
ax.set_xticks(xx); ax.set_xticklabels(MESES); ax.set_ylabel('CLP por pedido')
ax.set_title('Costo por pedido 2026 — in-house proyectado vs operador (tarifa + jefaturas integradas)', loc='left')
ax.legend(loc='lower left', frameon=False, fontsize=8); ax.grid(True, alpha=0.3)
save('13_costo_pedido')

# G14: Punto equilibrio LP (NUEVO v3)
fig, ax = plt.subplots(figsize=(11, 6))
volumenes = np.arange(50000, 600000, 10000)
colors_casos = [AZUL_OSC, AZUL_LOGO, AZUL]
labels_casos = ['Caso A: actual', 'Caso B: plan IA', 'Caso C: plan IA + eficiencia']
for i, caso in enumerate(casos_eq):
    c_in = (caso['F'] + caso['V'] * volumenes) / 1000  # MM CLP
    ax.plot(volumenes/1000, c_in/1000, label=f'In-house {labels_casos[i]}', color=colors_casos[i], linewidth=2)
# Operador (base)
c_op = (J_op + BENCH_OP_BASE * volumenes) / 1000
ax.plot(volumenes/1000, c_op/1000, label=f'Operador base ($2.000/ped)', color=NARANJA, linewidth=2.5, linestyle='--')
# Operador optimista
c_op_opt = (J_op + BENCH_OP_OPT * volumenes) / 1000
ax.plot(volumenes/1000, c_op_opt/1000, label=f'Operador OPT ($1.500/ped)', color=VERDE, linewidth=2.5, linestyle='--')
# Operador conservador
c_op_cons = (J_op + BENCH_OP_CONS * volumenes) / 1000
ax.plot(volumenes/1000, c_op_cons/1000, label=f'Operador CONS ($2.500/ped)', color=ROJO, linewidth=2.5, linestyle='--')
# Volumen actual y proyectado
ax.axvline(PEDIDOS_ANO/1000, color=GRIS_MED, linestyle=':', alpha=0.7, label=f'Volumen 2026: {PEDIDOS_ANO/1000:.0f}K')
ax.axvline(225, color=GRIS_MED, linestyle=':', alpha=0.5, label='Proyección 2027 optimista: 225K')
ax.set_xlabel('Pedidos al año (miles)'); ax.set_ylabel('Costo total (MM CLP)')
ax.set_title('Punto de equilibrio largo plazo — in-house vs operador (3 casos eficiencia + 3 tarifas)', loc='left')
ax.legend(frameon=False, loc='upper left', fontsize=8); ax.grid(True, alpha=0.3)
save('14_punto_equilibrio')

# G15: Costo por pedido 2027 por escenario (NUEVO v3)
fig, ax = plt.subplots(figsize=(11, 5))
nombres_e = [e['nombre'][:20] for e in esc_top]
x = np.arange(len(nombres_e))
w = 0.27
cxp_in = [e['cxp_inhouse_27'] for e in esc_top]
cxp_op_b = [e['cxp_op_27_base'] for e in esc_top]
cxp_op_o = [e['cxp_op_27_opt'] for e in esc_top]
b1 = ax.bar(x - w, cxp_in, w, label='In-house', color=AZUL_OSC, alpha=0.85)
b2 = ax.bar(x, cxp_op_b, w, label='Op. base ($2.000+jef)', color=NARANJA, alpha=0.85)
b3 = ax.bar(x + w, cxp_op_o, w, label='Op. opt ($1.500+jef)', color=VERDE, alpha=0.85)
for bars in [b1, b2, b3]:
    for b in bars: ax.text(b.get_x()+b.get_width()/2, b.get_height()+30, f'${b.get_height():,.0f}', ha='center', fontsize=7)
ax.set_xticks(x); ax.set_xticklabels(nombres_e, rotation=20, ha='right', fontsize=9)
ax.set_ylabel('CLP por pedido')
ax.set_title('Costo por pedido 2027 por escenario — in-house vs operador', loc='left')
ax.legend(frameon=False); ax.grid(True, axis='y', alpha=0.3)
save('15_cxp_2027')

# ===== DESCOMPOSICIÓN DEL EBITDA GAP POR PALANCA (NUEVO v5) =====
# EBITDA 2026 FCST: $268 MM (4,0%)
# EBITDA target 2027: 10% × V_target → necesitamos definir V_target
# Asumimos V_target = V26 × 1.08 = $7.192 MM (continuidad histórica) → EBITDA target = $719 MM
# Gap a cubrir: $719 - $268 = $451 MM
V_TARGET = V26 * (1 + CREC_25_26)  # $7.192 MM
EBITDA_TARGET = V_TARGET * 0.10  # $719 MM
EBITDA_GAP = EBITDA_TARGET - EBITDA26  # ~$451 MM

# Aporte estimado de cada palanca al EBITDA target
# Cada palanca tiene un rango (pesimista, realista, optimista) en MM CLP
palancas_ebitda = [
    # (categoría, palanca, pesimista, realista, optimista, factibilidad, dependencias)
    ('1. Venta',
     'Crecimiento de venta (mantiene MC% 27%, GAV escala parcial)',
     0, 35, 80,  # 0% crec, +8% (histórico), +15%
     'Realista: continuidad histórica +8% genera ~$35 MM EBITDA adicional vía absorción GAV',
     'Plan comercial + producto + marketing'),
    ('2. Margen Contribución',
     '2.1 Mejora compras + USD forward (+0.8-1.5 pp MC)',
     30, 67, 100,
     'Realista: 1pp de mejora sobre venta $6.700 MM = $67 MM',
     'Negociación volumen proveedores · Forward USD H1 2027'),
    ('2. Margen Contribución',
     '2.2 Renegociar tarifas Flex ML+Fala (+0.3-0.7 pp)',
     15, 35, 50,
     'Realista: 0,5 pp con volumen consolidado = $33 MM',
     'Forecast 6m por canal · negociación con couriers'),
    ('2. Margen Contribución',
     '2.3 Marketing digital eficiente (+0.5-1.0 pp)',
     20, 50, 70,
     'Realista: reasignar Mkt Performance hacia canales mejor ROI',
     'Análisis CAC y conversión por canal · plan IA RAW'),
    ('2. Margen Contribución',
     '2.4 Renegociar comisiones medios de pago web (+0.2-0.4 pp)',
     8, 20, 30,
     'Realista: Webpay/Transbank con volumen $51 MM web RM',
     'Negociación contractual con procesadores'),
    ('2. Margen Contribución',
     '2.5 Crecer share canales propios + B2B (+1.0-2.5 pp)',
     50, 130, 200,
     'Realista: cada +5% mix de marketplace a propio = ~0,5 pp MC',
     'Plan comercial B2B dedicado · inversión en web propios'),
    ('3. Eficiencia GAV',
     '3.1 Plan IA salidas (Jorgelis, Joselyn, Avila, Iris, Stipp + Camila V.)',
     45, 60, 75,
     'Confirmado: 6 cargos identificados con fechas. Reduce sueldos GAV en ~$76 MM/año',
     'Salidas plan IA H2 2026 y H1 2027 ejecutadas según plan'),
    ('3. Eficiencia GAV',
     '3.2 Reducción usuarios Odoo (35 → 15)',
     3, 5, 7,
     'Confirmado: $5,5 MM/año por simplificación de accesos',
     'Capacitación equipo · adopción Cerebro IA'),
    ('3. Eficiencia GAV',
     '3.3 Eliminación Multivende (jun 2026)',
     2, 2, 2,
     'Confirmado: $2 MM/año',
     'Integración alternativa lista'),
    ('3. Eficiencia GAV',
     '3.4 Audit consultora: duplicaciones (TradeMkt, asesorías ext, KAMs pequeños)',
     15, 30, 50,
     'Realista: $30 MM/año por consolidación cargos sobrevalorados o duplicados',
     'Decisión organizacional · reestructuración H1 2027'),
    ('3. Eficiencia GAV',
     '3.5 Reducción suscripciones SaaS no esenciales',
     5, 10, 15,
     'Auditoría licencias · cancelar duplicadas',
     'Audit IT Q4 2026'),
    ('3. Eficiencia GAV',
     '3.6 Eficiencia futura no comprometida (Camila Villalta y otros)',
     0, 22, 50,
     'Posible eficiencia adicional si se reorganiza Diseño Productos/Marketing. NO comprometida en plan base.',
     'Revisión organizacional Q3 2027 o posterior'),
    ('4. Operación logística',
     '4.1 Cambio a operador externo (tarifa BASE $2.000)',
     20, 52, 80,
     'Realista: $52 MM/año bajo tarifa base. Requiere negociación + indemnización $27 MM one-shot',
     'Negociación operador · validación servicio · plan transición H2 2026'),
    ('4. Operación logística',
     '4.1 alt Cambio a operador (tarifa OPT $1.500)',
     100, 142, 180,
     'Optimista: $142 MM/año. Depende de operador competitivo capturando volumen',
     'RFP a 3+ operadores · negociación agresiva'),
    ('4. Operación logística',
     '4.2 Eficiencia in-house (slotting, picking, rotación)',
     30, 60, 100,
     'Realista: $60 MM/año por mejora productividad operativa interna',
     'Plan IA Operaciones · capacitación · sistemas WMS'),
    ('5. Capital de Trabajo / Balance',
     '5.1 Reducir política inventario 5,5 → 4 m operativos',
     8, 15, 25,
     'Realista: libera $160 MM KT → reduce GNO ~$17 MM/año (no aporta EBITDA pero sí UN)',
     'Política inventario · plan slotting · forecast preciso'),
    ('5. Capital de Trabajo / Balance',
     '5.2 Reducir inventario obsoleto (10% → 5%)',
     5, 10, 15,
     'Realista: libera $80 MM KT por venta de obsoleto + no recompra',
     'Plan castigos · liquidación · prevención compras'),
    ('5. Capital de Trabajo / Balance',
     '5.3 Mejorar días CxC y CxP (negociación)',
     5, 10, 20,
     'Realista: cada 5 días mejora libera $30 MM KT → $3 MM GNO',
     'Negociación clientes B2B · proveedores · cobranza eficiente'),
    ('6. Estructura financiera',
     '6.1 Refinanciar deuda LP a tasa más baja',
     10, 30, 60,
     'Optimista: bajar tasa de 10,5% a 8% → ahorra $50 MM/año GNO',
     'Mercado de deuda · ratings · estructura legal'),
    ('6. Estructura financiera',
     '6.2 Acelerar amortización con caja excedente',
     5, 15, 30,
     'Pagar más deuda con caja libera GNO en años futuros',
     'Caja disponible · análisis costo oportunidad'),
    ('7. Costo Directo',
     '7.1 Negociación volumen proveedores (-1-2% costo)',
     30, 60, 100,
     'Realista: 1% sobre costo directo $3.106 MM = $31 MM. Hasta 2% con compromiso volumen',
     'Negociación proveedores top · plan compras consolidado'),
    ('7. Costo Directo',
     '7.2 Reducción mermas y devoluciones',
     5, 15, 30,
     'Realista: mejor control calidad + matching pedido-producto',
     'Plan IA SAC · auditoría calidad'),
    ('8. Otros',
     '8.1 Renegociación arriendos vencidos',
     5, 12, 25,
     'Renegociar arriendos al vencimiento (oficina + bodega)',
     'Cronograma vencimientos contratos'),
    ('8. Otros',
     '8.2 Optimización tributaria',
     5, 15, 30,
     'Estructura legal · beneficios fiscales · arrastres pérdidas',
     'Asesoría tributaria especializada'),
]

# Agrupar por categoría
import collections
cat_totales = collections.defaultdict(lambda: {'pes': 0, 'rea': 0, 'opt': 0, 'palancas': []})
for cat, p, pes, rea, opt, det, dep in palancas_ebitda:
    cat_totales[cat]['pes'] += pes
    cat_totales[cat]['rea'] += rea
    cat_totales[cat]['opt'] += opt
    cat_totales[cat]['palancas'].append((p, pes, rea, opt, det, dep))

# Total
TOTAL_PES = sum(c['pes'] for c in cat_totales.values())
TOTAL_REA = sum(c['rea'] for c in cat_totales.values())
TOTAL_OPT = sum(c['opt'] for c in cat_totales.values())

# Combinación recomendada (subset realista que cubra el gap sin sobre-contar overlaps)
# OJO: hay overlap entre palancas (ej: plan IA y eficiencia GAV se solapan, operador y eficiencia in-house se solapan)
# Por eso seleccionamos un SUBSET coherente que sume al gap
ruta_recomendada = [
    ('1. Venta', 'Crecimiento histórico +8%', 35),
    ('2. MC', '2.1 Compras + USD forward (+1 pp MC)', 67),
    ('2. MC', '2.3 Mkt digital eficiente (+0,7 pp MC)', 47),
    ('2. MC', '2.5 Crecer canales propios + B2B (+1 pp MC)', 67),
    ('3. GAV', '3.1 Plan IA salidas pleno', 60),
    ('3. GAV', '3.4 Audit consultora (duplicaciones moderadas)', 25),
    ('4. Operativo', '4.2 Eficiencia in-house (slotting/picking)', 60),
    ('7. CD', '7.1 Negociación proveedores (-1%)', 31),
]
TOTAL_RUTA = sum(p[2] for p in ruta_recomendada)
EBITDA_27_RUTA = EBITDA26 + TOTAL_RUTA * 1000  # convertir MM a miles
EBITDA_27_RUTA_PCT = EBITDA_27_RUTA / V_TARGET

# Ruta alternativa con operador (en lugar de eficiencia in-house)
ruta_alt = [
    ('1. Venta', 'Crecimiento histórico +8%', 35),
    ('2. MC', '2.1 Compras + USD forward (+1 pp)', 67),
    ('2. MC', '2.5 Crecer canales propios + B2B (+0,5 pp)', 33),
    ('3. GAV', '3.1 Plan IA salidas pleno', 60),
    ('3. GAV', '3.4 Audit consultora moderado', 25),
    ('4. Operativo', '4.1 Cambio operador base ($2.000)', 52),
    ('5. KT/Bal', '5.1 Pol inventario 4 m (impacto en UN, no EBITDA)', 0),
    ('7. CD', '7.1 Negociación proveedores (-1%)', 31),
]
TOTAL_RUTA_ALT = sum(p[2] for p in ruta_alt)
EBITDA_27_RUTA_ALT = EBITDA26 + TOTAL_RUTA_ALT * 1000
EBITDA_27_RUTA_ALT_PCT = EBITDA_27_RUTA_ALT / V_TARGET

# G16: Efecto operador por escenario (NUEVO v4)
fig, ax = plt.subplots(figsize=(12, 6))
nombres_e = [e['nombre'][:18] for e in esc_top]
x = np.arange(len(nombres_e))
w = 0.27
ebitda_in = [e['ebitda_pct']*100 for e in esc_top]
ebitda_op_b = [e['ebitda_pct_op_base']*100 for e in esc_top]
ebitda_op_o = [e['ebitda_pct_op_opt']*100 for e in esc_top]
b1 = ax.bar(x - w, ebitda_in, w, label='In-house (default)', color=AZUL_OSC, alpha=0.85)
b2 = ax.bar(x, ebitda_op_b, w, label='+ Operador base ($2.000+jef)', color=NARANJA, alpha=0.85)
b3 = ax.bar(x + w, ebitda_op_o, w, label='+ Operador opt ($1.500+jef)', color=VERDE, alpha=0.85)
ax.axhline(10, color=ROJO, linestyle='--', linewidth=2, label='Objetivo CEO: EBITDA 10%')
for bars in [b1, b2, b3]:
    for b in bars: ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.2, f'{b.get_height():.1f}%', ha='center', fontsize=7, fontweight='bold')
ax.set_xticks(x); ax.set_xticklabels(nombres_e, rotation=15, ha='right', fontsize=9)
ax.set_ylabel('EBITDA %')
ax.set_title('Efecto operador logístico — EBITDA % por escenario × estructura operativa', loc='left')
ax.legend(frameon=False, loc='upper left'); ax.grid(True, axis='y', alpha=0.3)
save('16_efecto_operador')

# G17: Waterfall del EBITDA — descomposición por palanca (RUTA RECOMENDADA)
fig, ax = plt.subplots(figsize=(13, 6))
cats = ['EBITDA\n2026 FCST'] + [p[1][:30] for p in ruta_recomendada] + ['EBITDA\n2027 ruta'] + ['Gap a\nobjetivo 10%']
vals = [EBITDA26/1000] + [p[2] for p in ruta_recomendada] + [EBITDA_27_RUTA/1000] + [EBITDA_TARGET/1000 - EBITDA_27_RUTA/1000]
# Cascada
cumul = EBITDA26/1000
bottoms = [0]
heights = [vals[0]]
colors_wf = [GRIS_OSC]
for i in range(1, len(vals)-2):
    bottoms.append(cumul)
    heights.append(vals[i])
    cumul += vals[i]
    colors_wf.append(VERDE)
# Total ruta
bottoms.append(0); heights.append(cumul); colors_wf.append(AZUL_OSC)
# Gap
gap_val = EBITDA_TARGET/1000 - cumul
bottoms.append(cumul); heights.append(gap_val); colors_wf.append(NARANJA if gap_val > 0 else VERDE)
ax.bar(range(len(cats)), heights, bottom=bottoms, color=colors_wf, alpha=0.85, width=0.7)
# Línea objetivo
ax.axhline(EBITDA_TARGET/1000, color=ROJO, linestyle='--', linewidth=2, label=f'Objetivo EBITDA 10% = ${EBITDA_TARGET/1000:,.0f} MM')
# Etiquetas
for i, (b, h) in enumerate(zip(bottoms, heights)):
    val_txt = f'${h:+,.0f}' if i not in [0, len(cats)-2, len(cats)-1] else f'${h:,.0f}'
    if i == len(cats)-1: val_txt = f'${h:+,.0f}'
    ax.text(i, b + h + 8, val_txt, ha='center', fontweight='bold', fontsize=9)
ax.set_xticks(range(len(cats))); ax.set_xticklabels(cats, rotation=20, ha='right', fontsize=8)
ax.set_ylabel('EBITDA (MM CLP)')
ax.set_title('Descomposición EBITDA 2026 → 2027 — Ruta multifactor recomendada (con +8% V histórico)', loc='left')
ax.legend(loc='upper left', frameon=False); ax.grid(True, axis='y', alpha=0.3)
save('17_waterfall_palancas')

# G18: Matriz palancas por categoría
fig, ax = plt.subplots(figsize=(12, 6))
cats_n = list(cat_totales.keys())
pes_v = [cat_totales[c]['pes'] for c in cats_n]
rea_v = [cat_totales[c]['rea'] for c in cats_n]
opt_v = [cat_totales[c]['opt'] for c in cats_n]
y = np.arange(len(cats_n))
h = 0.27
ax.barh(y-h, pes_v, h, label='Pesimista', color=ROJO, alpha=0.7)
ax.barh(y, rea_v, h, label='Realista', color=AZUL_OSC, alpha=0.85)
ax.barh(y+h, opt_v, h, label='Optimista', color=VERDE, alpha=0.7)
for i, c in enumerate(cats_n):
    ax.text(opt_v[i]+5, y[i]+h, f'${opt_v[i]:.0f}M', va='center', fontsize=8, color=VERDE_OSC)
    ax.text(rea_v[i]+5, y[i], f'${rea_v[i]:.0f}M', va='center', fontsize=9, fontweight='bold', color=AZUL_OSC)
    ax.text(pes_v[i]+5, y[i]-h, f'${pes_v[i]:.0f}M', va='center', fontsize=8, color=ROJO)
ax.set_yticks(y); ax.set_yticklabels([c[:35] for c in cats_n])
ax.set_xlabel('Aporte al EBITDA (MM CLP/año)')
ax.set_title('Aporte por categoría al EBITDA — escenarios pesimista, realista, optimista', loc='left')
ax.legend(frameon=False, loc='lower right'); ax.grid(True, axis='x', alpha=0.3)
save('18_matriz_categorias')

print(f"   ✓ 18 gráficos generados")


# =====================================================================
# GENERACIÓN DEL WORD
# =====================================================================
print("\n📄 Generando Word v3...")

doc = Document()
styles = doc.styles
norm = styles['Normal']
norm.font.name = 'Arial'; norm.font.size = Pt(10)
for section in doc.sections:
    section.left_margin = Cm(2); section.right_margin = Cm(2)
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)

def rgb(hex_str):
    s = hex_str.lstrip('#')
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.lstrip('#')); tc_pr.append(shd)

def set_cell_borders(cell, color='CBD5E1', size='4'):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_b = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), size); e.set(qn('w:color'), color)
        tc_b.append(e)
    tc_pr.append(tc_b)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.bold = True; run.font.name = 'Arial'
    sizes = {0: 22, 1: 18, 2: 14, 3: 12}
    colors_h = {0: AZUL_OSC, 1: AZUL_OSC, 2: AZUL_OSC, 3: GRIS_OSC}
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = rgb(colors_h.get(level, GRIS_OSC))
    p.paragraph_format.space_before = Pt([0, 24, 16, 10][min(level, 3)])
    p.paragraph_format.space_after = Pt([10, 10, 6, 4][min(level, 3)])
    return p

def add_callout(doc, text, color=AZUL_OSC):
    """Caja destacada para mensajes clave."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.text = ''
    set_cell_bg(cell, FONDO)
    # Borde lateral grueso del color
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_b = OxmlElement('w:tcBorders')
    left = OxmlElement('w:left'); left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24'); left.set(qn('w:color'), color.lstrip('#'))
    tc_b.append(left)
    for side in ['top', 'bottom', 'right']:
        e = OxmlElement(f'w:{side}'); e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4'); e.set(qn('w:color'), 'CBD5E1')
        tc_b.append(e)
    tc_pr.append(tc_b)
    # Padding
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'left', 'bottom', 'right']:
        m = OxmlElement(f'w:{side}'); m.set(qn('w:w'), '120'); m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tc_pr.append(tcMar)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Arial'; run.font.size = Pt(10); run.font.bold = True
    run.font.color.rgb = rgb(GRIS_OSC)
    cell.width = Inches(6.5)
    doc.add_paragraph()

def add_para(doc, text, bold=False, italic=False, size=10, color=GRIS_OSC, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.name = 'Arial'; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic; r.font.color.rgb = rgb(color)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify': p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    if p.runs:
        r = p.runs[0]; r.text = text
    else:
        r = p.add_run(text)
    r.font.name = 'Arial'; r.font.size = Pt(10); return p

def add_image(doc, path, width=6.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))

def add_table(doc, data, headers=None, col_widths=None, header_color=AZUL_OSC, alt=True):
    n = len(headers) if headers else len(data[0])
    t = doc.add_table(rows=1 if headers else 0, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if headers:
        h = t.rows[0].cells
        for i, hd in enumerate(headers):
            h[i].text = ''
            p = h[i].paragraphs[0]
            r = p.add_run(str(hd)); r.font.name = 'Arial'; r.font.size = Pt(9); r.font.bold = True
            r.font.color.rgb = rgb(BLANCO); set_cell_bg(h[i], header_color); set_cell_borders(h[i])
            h[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(data):
        rr = t.add_row()
        for ci, v in enumerate(row):
            c = rr.cells[ci]; c.text = ''
            p = c.paragraphs[0]; r = p.add_run(str(v))
            r.font.name = 'Arial'; r.font.size = Pt(9); r.font.color.rgb = rgb(GRIS_OSC)
            bg = (FONDO if alt and ri%2==0 else BLANCO)
            set_cell_bg(c, bg); set_cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def page_break(doc): doc.add_page_break()


# ===== PORTADA =====
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(LOGO, width=Inches(2.5))
doc.add_paragraph(); doc.add_paragraph()
add_para(doc, 'ANÁLISIS FINANCIERO', bold=True, size=14, color=AZUL_OSC, align='center')
add_para(doc, 'PROYECCIÓN 2027', bold=True, size=28, color=AZUL_OSC, align='center')
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Matriz extendida · Escenarios operador · Punto de equilibrio LP')
r.font.name = 'Arial'; r.font.size = Pt(14); r.font.italic = True; r.font.color.rgb = rgb(GRIS_MED)
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
add_para(doc, 'Preparado por', size=10, color=GRIS_MED, align='center')
add_para(doc, 'Andrés Browne', bold=True, size=12, color=GRIS_OSC, align='center')
add_para(doc, 'Gerencia Finanzas + Supply Chain', size=10, color=GRIS_MED, align='center')
doc.add_paragraph()
add_para(doc, 'Para presentación al Directorio', size=10, color=GRIS_MED, align='center')
add_para(doc, 'Mayo 2026 · Versión 3 (correcciones)', size=10, color=GRIS_MED, align='center')
page_break(doc)

# ===== ÍNDICE =====
add_heading(doc, 'ÍNDICE', level=0)
indice = [
    ('1', 'Resumen Ejecutivo'),
    ('2', 'Cierre Proyectado 2026'),
    ('3', 'Marco 2027 — Restricciones, supuestos, inventario y deuda'),
    ('4', 'Análisis Ácido Estático 2027'),
    ('5', 'Escenario 1: Margen Frontal y rentabilidad canales digitales'),
    ('6', 'Escenario 2: GAV y Costo Operativo — Análisis completo operador'),
    ('7', 'Matriz extendida de escenarios — MC 27% → 33%'),
    ('8', 'Sensibilidad ±20% sobre 2026'),
    ('9', 'Plan de Contingencia: ajustes desde lejano al core'),
    ('10', 'Recomendaciones de Factibilidad'),
    ('11', 'Anexos'),
]
for n, t in indice:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{n}.    '); r1.font.bold = True; r1.font.name = 'Arial'; r1.font.size = Pt(11)
    r1.font.color.rgb = rgb(AZUL_OSC)
    r2 = p.add_run(t); r2.font.name = 'Arial'; r2.font.size = Pt(11); r2.font.color.rgb = rgb(GRIS_OSC)
page_break(doc)

# ===== 1. RESUMEN EJECUTIVO =====
add_heading(doc, '1. RESUMEN EJECUTIVO', level=1)
add_para(doc, 'Contexto', bold=True, size=11, color=AZUL_OSC)
add_para(doc, ('El CEO estableció cinco objetivos prioritarios para 2027: D/E entre 2,5-3,5, '
               'sueldo CEO +10-15%, sueldo Erich vitalicio, retirar $100 MM (política 30%, UN ≥ '
               '$333 MM) y EBITDA del 10%. Este análisis cuantifica los requerimientos partiendo '
               f'del cierre proyectado 2026 y los supuestos confirmados. Cabe destacar que la '
               f'venta creció +{CREC_25_26*100:.1f}% entre 2025 y 2026, marcando un benchmark '
               'histórico de continuidad para evaluar la factibilidad de los escenarios 2027.'),
         align='justify')

doc.add_paragraph()
add_para(doc, 'Resultados clave', bold=True, size=11, color=AZUL_OSC)
resumen = [
    ['Métrica', '2026 FCST', '2027 Realista', '2027 Mejorado', '2027 Optimista'],
    ['Venta (MM CLP)', f'${V26/1000:,.0f}',
     f'${esc_top[1]["venta"]/1000:,.0f}',
     f'${esc_top[2]["venta"]/1000:,.0f}',
     f'${esc_top[3]["venta"]/1000:,.0f}'],
    ['Crecimiento V', '(referencia)', f'+{esc_top[1]["crec"]*100:.0f}%', f'+{esc_top[2]["crec"]*100:.0f}%', f'+{esc_top[3]["crec"]*100:.0f}%'],
    ['MC %', f'{MCp26*100:.1f}%', f'{esc_top[1]["mc"]*100:.0f}%', f'{esc_top[2]["mc"]*100:.0f}%', f'{esc_top[3]["mc"]*100:.0f}%'],
    ['EBITDA %', f'{EBITDA26/V26*100:.1f}%', f'{esc_top[1]["ebitda_pct"]*100:.1f}%',
     f'{esc_top[2]["ebitda_pct"]*100:.1f}%', f'{esc_top[3]["ebitda_pct"]*100:.1f}%'],
    ['UN (MM CLP)', f'${UN26/1000:,.0f}', f'${esc_top[1]["un"]/1000:,.0f}',
     f'${esc_top[2]["un"]/1000:,.0f}', f'${esc_top[3]["un"]/1000:,.0f}'],
    ['Cumple 5 restr. CEO', '—',
     '✓' if (esc_top[1]['cumple_ebitda'] and esc_top[1]['cumple_un'] and esc_top[1]['cumple_de']) else '✗',
     '✓' if (esc_top[2]['cumple_ebitda'] and esc_top[2]['cumple_un'] and esc_top[2]['cumple_de']) else '✗',
     '✓' if (esc_top[3]['cumple_ebitda'] and esc_top[3]['cumple_un'] and esc_top[3]['cumple_de']) else '✗'],
]
add_table(doc, resumen[1:], headers=resumen[0], col_widths=[1.8, 1.2, 1.4, 1.4, 1.4])

add_para(doc, 'Conclusiones top', bold=True, size=11, color=AZUL_OSC)
add_bullet(doc, (f'El escenario Realista (crecimiento histórico +{CREC_25_26*100:.0f}%, MC 29%, '
                 'eficiencia GAV 5%) NO cumple EBITDA 10%. Resultado proyectado: '
                 f'EBITDA {esc_top[1]["ebitda_pct"]*100:.1f}%. Esto significa que la continuidad '
                 'sin palancas adicionales no alcanza los objetivos del CEO.'))
add_bullet(doc, ('Para cumplir EBITDA 10% se requiere subir MC% sobre el histórico (al menos '
                 '31-33%) en combinación con eficiencia GAV. El escenario Mejorado (+12% V, '
                 'MC 31%) o el Optimista (+17% V, MC 33%) son los primeros que cumplen las '
                 '5 restricciones.'))
add_bullet(doc, ('Las palancas reales de mejora de MC% son: mejora de compras + USD forward · '
                 'reducción tarifas Flex · marketing digital eficiente · mejor comisión medios '
                 'de pago web · y especialmente mayor share de canales propios (web + B2B). '
                 'La renegociación con marketplaces NO es viable.'))
add_bullet(doc, (f'En el análisis de operador logístico, el operador externo conviene en 2026 '
                 f'sólo si la tarifa todo-incluida está en $1.500-2.000/ped. A $2.000 el ahorro '
                 'es modesto. A $2.500 pierde dinero. En largo plazo, in-house gana eficiencia '
                 'sobre volúmenes >300K pedidos/año.'))

doc.add_paragraph()
add_para(doc, 'Recomendaciones', bold=True, size=11, color=AZUL_OSC)
add_bullet(doc, ('Adoptar como meta corporativa el escenario Mejorado (+12% V, MC 31%, '
                 'eficiencia GAV 8%), con plan de stretch hacia Optimista si Q1 2027 muestra '
                 'tracción comercial > +15%.'))
add_bullet(doc, ('Bajar política inventario de 5,5 m a 4 m operativos. Libera ~$' +
                 f'{1.5*V26*0.466/12/1000:,.0f}' + ' MM en KT y reduce gasto financiero.'))
add_bullet(doc, ('Mantener operación logística in-house en 2026. Evaluar trimestralmente en 2027 '
                 'si la tarifa de operador llega a rango optimista ($1.500-1.800).'))

doc.add_paragraph()
add_callout(doc, 'CONCLUSIÓN CENTRAL: Los 5 objetivos del CEO se cumplen por orden de prioridad. '
                  'Los 3 primeros (D/E, sueldo CEO, sueldo Erich) son fáciles si la UN es positiva. '
                  'El 4° (UN ≥ $333 MM para retiros) requiere escenario Mejorado. El 5° (EBITDA 10%) '
                  'es el más exigente y demanda combinar múltiples palancas — NO se logra con una sola.',
            color=AZUL_OSC)
page_break(doc)

# ===== 2. CIERRE PROYECTADO 2026 =====
add_heading(doc, '2. CIERRE PROYECTADO 2026', level=1)
add_para(doc, ('Cifras del cierre proyectado 2026 (FCST mensual ene-dic) extraídas del archivo '
               'de Planificación Financiera. Son el ancla de todas las proyecciones 2027.'), align='justify')

add_heading(doc, '2.1 Estado de Resultados 2026', level=2)
pyl_data = [
    ['Concepto', 'MM CLP', '% Venta'],
    ['Ingresos por Ventas', f'${V26/1000:,.0f}', '100,0%'],
    ['Costo Directo', f'-${CD26/1000:,.0f}', f'-{CD26/V26*100:.1f}%'],
    ['Margen Frontal', f'${MF26/1000:,.0f}', f'{MFp26*100:.1f}%'],
    ['Otros Costos Explot. (comisión + flete + Mkt Perf)', f'-${OCE26/1000:,.0f}', f'-{OCE26/V26*100:.1f}%'],
    ['Margen Contribución', f'${MC26/1000:,.0f}', f'{MCp26*100:.1f}%'],
    ['GAV', f'-${GAV26/1000:,.0f}', f'-{GAV26/V26*100:.1f}%'],
    ['EBIT', f'${EBIT26/1000:,.0f}', f'{EBIT26/V26*100:.2f}%'],
    ['EBITDA', f'${EBITDA26/1000:,.0f}', f'{EBITDA26/V26*100:.2f}%'],
    ['GNO (intereses)', f'-${GNO26/1000:,.0f}', f'-{GNO26/V26*100:.1f}%'],
    ['Utilidad Neta', f'${UN26/1000:,.0f}', f'{UN26/V26*100:.2f}%'],
]
add_table(doc, pyl_data[1:], headers=pyl_data[0], col_widths=[3.5, 1.5, 1.2])
add_image(doc, os.path.join(CHARTS_DIR, '03_waterfall.png'), width=6.3)
add_para(doc, 'Figura 1. Cascada P&L 2026', italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '2.2 Crecimiento histórico y estacionalidad', level=2)
crec_data = [
    ['Año', 'Venta (MM)', 'MC %', 'Crecimiento'],
    ['2024', f'$5.978', f'28,3%', '—'],
    ['2025', f'${V25/1000:,.0f}', f'{MCp25*100:.1f}%', '+3,1%'],
    ['2026 FCST', f'${V26/1000:,.0f}', f'{MCp26*100:.1f}%', f'+{CREC_25_26*100:.1f}%'],
]
add_table(doc, crec_data[1:], headers=crec_data[0], col_widths=[1.5, 1.5, 1.5, 1.5])

add_para(doc, (f'El crecimiento histórico 2025→2026 fue +{CREC_25_26*100:.1f}%. Este es el '
               'benchmark realista de continuidad para definir escenarios 2027.'), align='justify', italic=True)

add_image(doc, os.path.join(CHARTS_DIR, '01_venta_mensual.png'), width=6.3)
add_para(doc, 'Figura 2. Venta, MC y EBIT mensuales 2026', italic=True, size=9, color=GRIS_MED, align='center')
add_image(doc, os.path.join(CHARTS_DIR, '09_estacionalidad.png'), width=6.3)
add_para(doc, 'Figura 3. Estacionalidad — peso porcentual de cada mes', italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '2.3 Estructura GAV', level=2)
add_image(doc, os.path.join(CHARTS_DIR, '02_gav_estructura.png'), width=5.5)
add_para(doc, 'Figura 4. Composición del GAV 2026', italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '2.4 Existencias mensuales — saldos reales planilla', level=2)
add_para(doc, ('Saldos reales al cierre de cada mes 2026 (datos de la planilla, sin aplicar '
               'supuestos). Estos valores corresponden al inventario en balance contable '
               '(bodega Chile).'), align='justify')
exist_data = [['Mes', 'Existencias (MM)', 'Meses inventario']]
for i, m in enumerate(MESES):
    exist_data.append([m, f'${existencias_m[i]/1000:,.0f}', f'{meses_exist_m[i]:.2f}'])
exist_data.append(['Promedio', f'${EXIST_PROM/1000:,.0f}', f'{MESES_EXIST_PROM:.2f}'])
add_table(doc, exist_data[1:], headers=exist_data[0], col_widths=[1.0, 2.0, 2.0])
add_image(doc, os.path.join(CHARTS_DIR, '10_existencias.png'), width=6.3)
add_para(doc, 'Figura 5. Existencias mensuales 2026 con líneas política (5,5 m) y recomendación (4 m)',
         italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '2.5 Balance, deuda y capital de trabajo', level=2)
bal_data = [
    ['Concepto', 'Valor (MM CLP)'],
    ['Patrimonio Neto (dic 2026)', f'${PATRIM26/1000:,.0f}'],
    ['Deuda Financiera Total (dic 2026)', f'${DEUDA_TOTAL_DIC/1000:,.0f}'],
    ['  - COMEX (rotatorio con compras)', f'${COMEX_DIC/1000:,.0f}'],
    ['  - Crédito en pesos (amortizando)', f'${CREDITO_PESOS_DIC/1000:,.0f}'],
    ['Ratio Deuda/Patrimonio (D/E)', f'{DE26:.2f}x ✓ en rango bajo riesgo'],
    ['Capital de Trabajo Neto (dic)', f'${kt_neto_m[-1]/1000:,.0f}'],
    ['Tasa de interés efectiva sobre deuda', f'{TASA_INT_EFEC*100:.1f}% anual'],
]
add_table(doc, bal_data[1:], headers=bal_data[0], col_widths=[4.0, 2.5])
page_break(doc)

# ===== 3. MARCO 2027 =====
add_heading(doc, '3. MARCO 2027 — RESTRICCIONES, SUPUESTOS, INVENTARIO Y DEUDA', level=1)

add_heading(doc, '3.1 Restricciones del CEO', level=2)
ceo_data = [
    ['#', 'Objetivo', 'Cálculo / Implicancia'],
    ['1', 'D/E en rango bajo riesgo (2,5-3,5)', f'Con deuda ~$2.000 MM, patrimonio entre $570-$800 MM'],
    ['2', 'Sueldo CEO +10-15%', 'Martín Novoa: $9,87 MM/m × 1,12 = ~$11,05 MM/m'],
    ['3', 'Sueldo Erich vitalicio', 'Mantener $5,21 MM/m'],
    ['4', 'Retirar $100 MM (política 30%)', 'UN 2027 ≥ $333 MM'],
    ['5', 'EBITDA 10%', 'EBITDA 2027 = 10% × Venta'],
]
add_table(doc, ceo_data[1:], headers=ceo_data[0], col_widths=[0.5, 2.5, 3.5])

add_heading(doc, '3.2 Supuestos confirmados', level=2)
sup_data = [
    ['Supuesto', 'Valor'],
    ['Tasa impuesto', '27%'],
    ['Tasa interés efectiva sobre deuda', f'{TASA_INT_EFEC*100:.1f}%'],
    ['Ajuste UF gastos varios', '4%'],
    ['Ajuste salarial general (incluye CEO)', '12%'],
    ['Plan IA salidas', '5 cargos plan IA + Camila Villalta'],
    ['Control de Gestión (Sub-gerencia Finanzas)', 'SE MANTIENE en 2027 (no sale)'],
    ['Marketing Branding', '3% de la venta'],
    ['Inventario política operativa', '5,5 meses (3,5+1+1) en versión CEO · 4 m recomendado'],
    ['Inventario obsoleto', '10% recurrente'],
    ['Costo por pedido in-house actual (KPI app)', '$2.812'],
    ['Crecimiento histórico 2025→2026', f'+{CREC_25_26*100:.1f}%'],
]
add_table(doc, sup_data[1:], headers=sup_data[0], col_widths=[3.5, 3.0])

add_heading(doc, '3.3 Tratamiento del inventario y relación con KT', level=2)
add_para(doc, ('La sección 2.4 mostró saldos reales 2026: promedio 4,4 meses, peak feb 5,9, valle '
               'nov 2,4. La política operativa del CEO para 2027 es 5,5 meses (3,5 bodega + '
               '1 tránsito + 1 producción) con 10% obsoleto recurrente.'), align='justify')

costo_venta_27_base = esc_top[1]['venta'] * (CD26/V26)
inv_bodega_27 = costo_venta_27_base * 3.5 / 12
inv_obsoleto = inv_bodega_27 * 0.10 / 0.90
inv_data = [
    ['Componente', 'Meses', 'MM CLP (2027 Realista)', 'Aparece en balance'],
    ['Bodega Chile operativo', '3,5', f'${inv_bodega_27/1000:,.0f}', 'Sí'],
    ['Tránsito (mar/aire)', '1,0', f'${costo_venta_27_base/12/1000:,.0f}', 'Anticipo proveedores'],
    ['Producción China', '1,0', f'${costo_venta_27_base/12/1000:,.0f}', 'Anticipo proveedores'],
    ['Obsoleto recurrente (10% bodega)', '—', f'${inv_obsoleto/1000:,.0f}', 'Sí'],
    ['Existencias en balance (bodega + obs)', '~3,9', f'${(inv_bodega_27+inv_obsoleto)/1000:,.0f}', 'Sí'],
]
add_table(doc, inv_data[1:], headers=inv_data[0], col_widths=[2.5, 0.8, 2.0, 1.7])

add_para(doc, ('Recomendación crítica: pasar la política bodega de 3,5 m a 2,5 m operativos ' +
               '(política total 4 m en lugar de 5,5 m). Libera ~$' +
               f'{1.0*costo_venta_27_base/12/1000:,.0f}' +
               ' MM en KT, reduce el COMEX necesario y mejora el D/E. Coherente con '
               'el comportamiento real 2026 (promedio 4,4 m).'),
         align='justify', italic=True)

add_heading(doc, '3.4 Estructura de Deuda', level=2)
add_para(doc, ('La deuda financiera ($' + f'{DEUDA_TOTAL_DIC/1000:,.0f}' + ' MM dic 2026) se '
               'compone de dos instrumentos con comportamientos opuestos:'), align='justify')
deuda_data = [
    ['Tipo', 'Saldo ene-26', 'Saldo dic-26', 'Comportamiento', 'Tasa estim.'],
    ['Crédito en pesos', f'${credito_pesos_m[0]/1000:,.0f}',
     f'${CREDITO_PESOS_DIC/1000:,.0f}', 'AMORTIZANDO (cuota mensual)', '~8,5% anual'],
    ['COMEX (multibanco)', f'${comex_m[0]/1000:,.0f}',
     f'${COMEX_DIC/1000:,.0f}', 'ROTATORIO (crece con compras)', '~5-6% anual'],
    ['Total', f'${deuda_total_m[0]/1000:,.0f}',
     f'${DEUDA_TOTAL_DIC/1000:,.0f}', '', f'{TASA_INT_EFEC*100:.1f}% mezcla'],
]
add_table(doc, deuda_data[1:], headers=deuda_data[0], col_widths=[1.7, 1.2, 1.2, 1.6, 1.2])
add_image(doc, os.path.join(CHARTS_DIR, '11_deuda.png'), width=6.3)
add_para(doc, 'Figura 6. Estructura de Deuda 2026', italic=True, size=9, color=GRIS_MED, align='center')
page_break(doc)

# ===== 4. ANÁLISIS ÁCIDO ESTÁTICO =====
add_heading(doc, '4. ANÁLISIS ÁCIDO ESTÁTICO 2027', level=1)
add_para(doc, ('El escenario ácido aplica las restricciones CEO sin asumir mejoras: MC% se '
               'mantiene en 27%, plan IA sin Gabriela (se mantiene), UF 4%, ajuste salarial 12%.'), align='justify')

add_heading(doc, '4.1 GAV 2027 proyectado', level=2)
gav27 = calc_gav_2027(V_BASE_EBITDA_27)
gav_t = [
    ['Concepto', '2026', '2027 ácido base', 'Δ'],
    ['Sueldos GAV', f'${SUELDOS26/1000:,.0f}', f'${gav27["sueldos"]/1000:,.0f}', f'{(gav27["sueldos"]/SUELDOS26-1)*100:+.1f}%'],
    ['Oficina/Arriendos', f'${OFICINA26/1000:,.0f}', f'${gav27["oficina"]/1000:,.0f}', '+4,0%'],
    ['Suscripciones (con Claude IA)', f'${SUSCRIPCIONES26/1000:,.0f}', f'${gav27["suscripciones"]/1000:,.0f}',
     f'{(gav27["suscripciones"]/SUSCRIPCIONES26-1)*100:+.0f}%'],
    ['Marketing Branding (3%V)', f'${MKT_BRANDING26/1000:,.0f}', f'${gav27["mkt_branding"]/1000:,.0f}', ''],
    ['Depreciación + Otros', f'${(GAV26-SUELDOS26-OFICINA26-SUSCRIPCIONES26-MKT_BRANDING26)/1000:,.0f}',
     f'${(gav27["depreciacion"]+gav27["otros"])/1000:,.0f}', ''],
    ['Total GAV', f'${GAV26/1000:,.0f}', f'${gav27["total"]/1000:,.0f}', f'{(gav27["total"]/GAV26-1)*100:+.1f}%'],
]
add_table(doc, gav_t[1:], headers=gav_t[0], col_widths=[2.3, 1.3, 1.5, 1.0])

add_heading(doc, '4.2 Cuello de botella entre las restricciones CEO', level=2)
v_un_27 = venta_para_un(333000)
binding = [
    ['Restricción', 'Venta requerida (MM)', 'Crecimiento vs 2026'],
    ['UN ≥ $333 MM (retiros $100 MM)', f'${v_un_27/1000:,.0f}', f'+{(v_un_27/V26-1)*100:.0f}%'],
    ['EBITDA = 10% ⭐ binding', f'${V_BASE_EBITDA_27/1000:,.0f}', f'+{(V_BASE_EBITDA_27/V26-1)*100:.0f}%'],
]
add_table(doc, binding[1:], headers=binding[0], col_widths=[2.5, 2.0, 2.0])

add_para(doc, (f'La restricción binding es EBITDA 10%. Sin palancas adicionales exige '
               f'crecimiento +{(V_BASE_EBITDA_27/V26-1)*100:.0f}% sobre 2026 — muy lejos del '
               f'histórico (+{CREC_25_26*100:.1f}%). En la Sección 7 (Matriz extendida) se '
               'explora cómo combinar palancas para reducir este crecimiento requerido a '
               'niveles factibles.'), align='justify')

add_heading(doc, '4.3 Análisis de Deuda y D/E por escenario', level=2)
add_para(doc, ('Una de las restricciones prioritarias del CEO es mantener el D/E en rango de '
               'bajo riesgo (2,5-3,5). Este ratio cambia en 2027 según: (a) la utilidad neta '
               'del escenario, (b) los retiros realizados (política 30%, máximo $100 MM/año) y '
               '(c) la evolución de la deuda (COMEX rotatorio + crédito en pesos amortizando).'), align='justify')

add_para(doc, 'Cálculo del D/E proyectado 2027 por escenario:', bold=True, size=10, color=GRIS_OSC)
add_para(doc, 'Patrimonio 2027 = Patrimonio 2026 + Utilidad Neta 2027 − Retiros 2027',
         size=10, color=AZUL_OSC, align='center', bold=True)
add_para(doc, ('Donde retiros 2027 = min(UN × 30%, $100 MM). Asumimos deuda 2027 estable en '
               f'~${DEUDA_TOTAL_DIC/1000:,.0f} MM (COMEX se mantiene con compras + crédito en '
               'pesos amortizando se compensa).'), align='justify', size=9, italic=True)

de_data = [
    ['Escenario', 'UN (MM)', 'Retiros (MM)', 'Patrimonio fin 2027 (MM)', 'D/E proyectado', '¿En rango 2,5-3,5?'],
]
for e in esc_top:
    patrim = PATRIM26 + e['un'] - min(e['un']*0.30, 100000)
    de = DEUDA_TOTAL_DIC / patrim if patrim > 0 else 999
    en_rango = '✓ Sí' if 2.5 <= de <= 3.5 else ('⚠ Bajo el rango (más conservador)' if de < 2.5 else '✗ Excede el rango')
    de_data.append([
        e['nombre'][:25],
        f'${e["un"]/1000:,.0f}',
        f'${min(e["un"]*0.30, 100000)/1000:,.0f}',
        f'${patrim/1000:,.0f}',
        f'{de:.2f}x',
        en_rango,
    ])
add_table(doc, de_data[1:], headers=de_data[0], col_widths=[1.7, 0.8, 0.9, 1.4, 1.0, 1.7])

add_para(doc, ('Lectura del D/E proyectado:'), bold=True, size=10, color=AZUL_OSC)
add_bullet(doc, ('En el escenario Status quo (UN negativa o muy baja), el patrimonio cae y el '
                 'D/E supera 3,5 — sale del rango aceptable.'))
add_bullet(doc, ('En el escenario Realista (UN ~$280 MM), el patrimonio crece y el D/E baja '
                 'a ~2,4x. Está por DEBAJO del rango pero es bueno (menor riesgo).'))
add_bullet(doc, ('En escenarios Mejorado y superiores, el D/E baja aún más (1,7-2,0x). El CEO '
                 'pidió "bajo riesgo": estar bajo el rango no es problema, es mejor.'))
add_bullet(doc, ('Implicancia estratégica: el cumplimiento del D/E está condicionado a tener '
                 'UN positiva razonable. Cualquier escenario que cumpla el objetivo UN ≥ $333 MM '
                 'automáticamente cumple el D/E con holgura.'))

add_heading(doc, '4.4 Estructura de deuda 2027 — proyección', level=2)
add_para(doc, ('La deuda total 2027 dependerá del comportamiento de los dos instrumentos:'), align='justify')

deuda_27_data = [
    ['Instrumento', 'Saldo dic 2026', 'Tendencia 2027', 'Saldo estimado dic 2027'],
    ['Crédito en pesos', f'${CREDITO_PESOS_DIC/1000:,.0f} MM',
     'Amortiza ~$10 MM/mes (cuota fija)',
     f'~${max(0, CREDITO_PESOS_DIC-120000)/1000:,.0f} MM (-32%)'],
    ['COMEX (rotatorio)', f'${COMEX_DIC/1000:,.0f} MM',
     'Crece con compras de inventario',
     'Estable si política inventario 4m · Sube ~10% si crece venta +12%'],
    ['Total estimado', f'${DEUDA_TOTAL_DIC/1000:,.0f} MM',
     '', f'~$1.900-2.100 MM (escenarios moderados)'],
]
add_table(doc, deuda_27_data[1:], headers=deuda_27_data[0], col_widths=[1.8, 1.4, 2.0, 1.7])

add_para(doc, ('Palancas adicionales sobre el D/E (no necesariamente vía EBITDA):'), bold=True, size=10, color=AZUL_OSC)
add_bullet(doc, ('Reducir política inventario 5,5 → 4m libera ~$160 MM en COMEX → baja deuda → '
                 'mejora D/E directamente.'))
add_bullet(doc, ('Refinanciar deuda a tasa más baja (de 10,5% a 8%) no afecta D/E pero ahorra '
                 'gastos financieros (~$50 MM/año) que impactan en UN.'))
add_bullet(doc, ('Acelerar amortización con caja excedente baja deuda absoluta → mejora D/E '
                 'al alza (más conservador).'))

doc.add_paragraph()
add_callout(doc, 'CONCLUSIÓN D/E: el cumplimiento del rango 2,5-3,5 está condicionado a tener UN positiva razonable. '
                  'Cualquier escenario que cumpla la restricción UN ≥ $333 MM (Escenario Mejorado o superior) '
                  'cumple el D/E automáticamente con holgura (queda en 1,7-2,2x). En escenarios débiles '
                  '(Status quo o Realista bajo), el D/E queda al filo o se sale del rango por arriba.',
            color=NARANJA)
page_break(doc)

# ===== 5. ESCENARIO 1 MARGEN =====
add_heading(doc, '5. ESCENARIO 1: MARGEN FRONTAL Y RENTABILIDAD CANALES DIGITALES', level=1)
add_para(doc, ('Palancas comerciales y de compras que mejoran el margen de contribución. La '
               'renegociación con marketplaces NO se incluye (no es viable). Las palancas '
               'reales son: mejora margen directo, tarifas Flex, marketing digital eficiente, '
               'medios de pago web y mix de canales propios + B2B.'), align='justify')

add_heading(doc, '5.1 Estacionalidad del margen — análisis mensual 2026', level=2)
est_data = [['Mes', 'Venta (MM)', 'MF %', 'Comisión %', 'MC %']]
for i, m in enumerate(MESES):
    mfp = mf_m[i]/ventas_m[i]*100 if ventas_m[i] else 0
    cmp = com_m[i]/ventas_m[i]*100 if ventas_m[i] else 0
    mcp = mc_m[i]/ventas_m[i]*100 if ventas_m[i] else 0
    est_data.append([m, f'${ventas_m[i]/1000:,.0f}', f'{mfp:.1f}%', f'{cmp:.1f}%', f'{mcp:.1f}%'])
est_data.append(['Promedio', f'${V26/12000:,.0f}', f'{MFp26*100:.1f}%', f'{sum(com_m)/V26*100:.1f}%', f'{MCp26*100:.1f}%'])
add_table(doc, est_data[1:], headers=est_data[0], col_widths=[0.8, 1.5, 1.0, 1.2, 1.0])
add_image(doc, os.path.join(CHARTS_DIR, '12_margen_mensual.png'), width=6.3)
add_para(doc, 'Figura 7. Estacionalidad del margen 2026', italic=True, size=9, color=GRIS_MED, align='center')

add_para(doc, ('Los meses peak (Jun, Oct, Dic) tienen comisión ~24% s/venta vs 18% en meses '
               'tranquilos. Esto refleja el peso de marketplaces en el peak comercial.'), align='justify')

add_heading(doc, '5.2 Palancas reales para mejorar MC%', level=2)
palancas = [
    ['#', 'Palanca', 'Impacto MC'],
    ['1', 'Mejora compras + USD forward (margen directo)', '+0,8 a +1,5 pp'],
    ['2', 'Tarifas Flex (Falabella y ML)', '+0,3 a +0,7 pp'],
    ['3', 'Marketing digital eficiente (reasignar gasto a canales con mejor ROI)', '+0,5 a +1,0 pp'],
    ['4', 'Renegociar comisiones medios de pago web (Webpay)', '+0,2 a +0,4 pp'],
    ['5', 'Mayor share canales propios y B2B (MC 65-70%)', '+1,0 a +2,5 pp'],
    ['TOTAL', 'Combinación conservadora', '+3 a +6 pp (MC% 30-33%)'],
]
add_table(doc, palancas[1:], headers=palancas[0], col_widths=[0.5, 4.5, 1.5])
page_break(doc)

# ===== 6. ESCENARIO 2: GAV Y COSTO OPERATIVO =====
add_heading(doc, '6. ESCENARIO 2: GAV Y COSTO OPERATIVO — ANÁLISIS COMPLETO OPERADOR', level=1)

add_heading(doc, '6.1 Distribución del payroll', level=2)
add_heading(doc, 'A. Por nivel jerárquico', level=3)
agg_nivel = df_pay_real.groupby('NIVEL')['TH'].agg(['sum', 'count']).reset_index().sort_values('NIVEL')
total_real = df_pay_real['TH'].sum()
nivel_data = [['Nivel', '# personas', 'Costo mensual', '% total', 'Anualizado']]
for _, row in agg_nivel.iterrows():
    if row['sum'] > 0:
        nivel_data.append([str(row['NIVEL']), f'{int(row["count"])}',
                          f'${row["sum"]/1e6:,.2f} MM', f'{row["sum"]/total_real*100:.1f}%',
                          f'${row["sum"]*12/1e6:.0f} MM'])
nivel_data.append(['TOTAL', f'{len(df_pay_real)}', f'${total_real/1e6:.2f} MM',
                  '100,0%', f'${total_real*12/1e6:.0f} MM'])
add_table(doc, nivel_data[1:], headers=nivel_data[0], col_widths=[2.5, 0.8, 1.2, 0.8, 1.2])

add_para(doc, ('Las 5 gerencias confirmadas (Martín Novoa CEO, Erich Co-Founder, Andrés Browne '
               'GAF, Nicolás Vásquez Comercial, Sebastián Guzmán Productos) están separadas en '
               'niveles 1.1-1.5. Gabriela Pastran (Control de Gestión) se clasifica en nivel '
               '2.1 como Sub-gerencia Finanzas y se mantiene en 2027.'),
         align='justify', size=9, italic=True)

add_heading(doc, 'B. Por área y sub-área', level=3)
agg_area = df_pay_real.groupby(['AREA', 'SUBAREA'])['TH'].agg(['sum', 'count']).reset_index().sort_values('sum', ascending=False)
area_data = [['Área', 'Sub-área', 'n', 'Mensual', '% total']]
for _, row in agg_area.iterrows():
    if row['sum'] > 0:
        area_data.append([str(row['AREA'])[:15], str(row['SUBAREA'])[:25],
                          f'{int(row["count"])}', f'${row["sum"]/1e6:,.2f} MM',
                          f'{row["sum"]/total_real*100:.1f}%'])
add_table(doc, area_data[1:], headers=area_data[0], col_widths=[1.5, 1.8, 0.6, 1.2, 0.8])

add_heading(doc, '6.2 Costo por pedido — proyección con salidas plan IA', level=2)
add_para(doc, (f'El costo por pedido actual (KPI app) es ${COSTO_X_PED_ACTUAL:,}. Las salidas '
               'del plan IA aplicables a operaciones (Jorgelis Flores Log. Inv. 31-jul, Iris '
               'Yori Facturación 15-ago, Fernanda Stipp Facturación 15-sep) reducen el costo '
               f'a partir de su salida. Costo proyectado dic 2026 (full impact): ${COSTO_X_PED_POST_SALIDAS:,.0f}/ped.'),
         align='justify')
add_image(doc, os.path.join(CHARTS_DIR, '13_costo_pedido.png'), width=6.3)
add_para(doc, 'Figura 8. Costo por pedido 2026 — in-house proyectado vs benchmarks operador (tarifa + jefaturas integradas)',
         italic=True, size=9, color=GRIS_MED, align='center')

add_para(doc, ('Modelo operador externo: el operador opera con su propia gente y arriendos, '
               'cobra una tarifa todo-incluida por pedido ($1.500-2.500). Mantenemos integradas '
               'a sus operaciones a las 3 jefaturas operativas (Max Bellolio Post Venta, Gerardo '
               'Ortega Logística, Yohana Grisman Facturación) MÁS el Gerente de Operaciones '
               '(Andrés Browne, 50% de su sueldo asignado a Operaciones — el otro 50% queda en '
               f'Finanzas). Total costo retenido in-house: ${COSTO_JEFATURAS_INTEG/1000:,.0f} '
               f'MM/año = ${COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO:,.0f}/pedido sobre volumen actual. '
               'Esto eleva significativamente el costo retenido respecto a un modelo sin supervisión '
               'gerencial de operaciones.'),
         align='justify')

add_heading(doc, 'Escenario A — Cambio a operador agosto 2026', level=3)
add_para(doc, (f'Cambio efectivo 1-ago-2026. Las {N_OP_SALEN} personas operativas (sin jefaturas) '
               f'salen con indemnización legal Chile. Indemnización total: ${INDEMN_OP_SALEN/1e6:,.0f} '
               'MM. Comparación año completo:'), align='justify')
esc_a_data = [['Bench operador', 'In-house ene-jul', 'Operador ago-dic',
               'Indemnización', 'Total con cambio', 'Total sin cambio', 'Δ neto 2026']]
for ec in cambio_2026:
    esc_a_data.append([
        f'${ec["bench"]:,} ({ec["label"]})',
        f'${ec["costo_in_ene_jul"]/1000:,.0f}',
        f'${ec["costo_op_ago_dic"]/1000:,.0f}',
        f'${ec["indem"]/1000:,.0f}',
        f'${ec["total_con_cambio"]/1000:,.0f}',
        f'${ec["total_sin_cambio"]/1000:,.0f}',
        f'${ec["delta_neto"]/1000:+,.0f}',
    ])
add_table(doc, esc_a_data[1:], headers=esc_a_data[0], col_widths=[1.5, 1.0, 1.0, 0.9, 1.1, 1.0, 0.9])

opt_d = cambio_2026[0]['delta_neto']/1000
base_d = cambio_2026[1]['delta_neto']/1000
cons_d = cambio_2026[2]['delta_neto']/1000
add_para(doc, (f'Lectura: el cambio AHORRA ${opt_d:,.0f} MM bajo operador optimista ($1.500/ped), '
               f'genera diferencial neutro o leve de ${base_d:+,.0f} MM bajo base ($2.000/ped), '
               f'y PIERDE ${abs(cons_d):,.0f} MM bajo conservador ($2.500/ped). El benchmark '
               'pivote es $1.700-1.800/ped: bajo eso el cambio claramente conviene.'),
         align='justify')

add_heading(doc, 'Escenario B — 2027 por escenario combinado', level=3)
esc_b_data = [['Escenario', 'Venta (MM)', 'Pedidos', 'Cxp in-house', 'Cxp operador base', 'Cxp operador opt',
               'Decisión']]
for e in esc_top:
    decision = ('In-house' if e['cxp_inhouse_27'] < e['cxp_op_27_base'] else 'Operador base'
                if e['cxp_op_27_base'] < e['cxp_op_27_opt'] else 'Operador opt')
    esc_b_data.append([
        e['nombre'][:18], f'${e["venta"]/1000:,.0f}', f'{e["pedidos"]/1000:,.0f}K',
        f'${e["cxp_inhouse_27"]:,.0f}', f'${e["cxp_op_27_base"]:,.0f}',
        f'${e["cxp_op_27_opt"]:,.0f}', decision,
    ])
add_table(doc, esc_b_data[1:], headers=esc_b_data[0], col_widths=[1.5, 0.9, 0.7, 1.0, 1.1, 1.1, 1.0])

add_image(doc, os.path.join(CHARTS_DIR, '15_cxp_2027.png'), width=6.3)
add_para(doc, 'Figura 9. Costo por pedido 2027 por escenario — in-house vs operador base y optimista',
         italic=True, size=9, color=GRIS_MED, align='center')

add_para(doc, ('A mayor escenario de venta, el costo unitario in-house baja por mejor absorción '
               'de fijos. El operador optimista ($1.500+jef) es competitivo en todos los '
               'escenarios. El operador base ($2.000+jef) compite parejo con in-house en '
               'escenarios medios.'), align='justify')

add_heading(doc, 'Escenario C — ¿Cuándo conviene cambiar a operador externo? (Análisis directo)', level=3)

add_para(doc, ('Las decisiones de mantener in-house o cambiar a operador se toman cuando, '
               'según el escenario de venta proyectada, el costo TOTAL por pedido sea menor '
               'en una u otra alternativa. La siguiente tabla muestra esa comparación directa '
               'por cada escenario 2027 considerado.'), align='justify')

add_heading(doc, 'C.1 Proyección de pedidos 2027 por escenario', level=3)
add_para(doc, (f'Asumiendo un AOV (ticket promedio) constante de ${AOV:,} CLP, los pedidos '
               '2027 escalan proporcionalmente con la venta:'), align='justify')

ped_data = [['Escenario', 'Crec V', 'Venta 2027 (MM)', 'Pedidos 2027 (miles)', 'Δ vs 2026 (180K)']]
for e in esc_top:
    delta = (e["pedidos"]/PEDIDOS_ANO - 1) * 100
    ped_data.append([
        e['nombre'][:25], f'+{e["crec"]*100:.0f}%',
        f'${e["venta"]/1000:,.0f}', f'{e["pedidos"]/1000:,.0f}K',
        f'+{delta:.0f}% pedidos',
    ])
add_table(doc, ped_data[1:], headers=ped_data[0], col_widths=[2.0, 0.8, 1.5, 1.4, 1.3])

add_heading(doc, 'C.2 Comparación in-house vs operador por escenario', level=3)
add_para(doc, ('Para cada escenario 2027 calculamos el costo POR PEDIDO bajo 3 alternativas:'),
         align='justify')
add_bullet(doc, (f'IN-HOUSE con plan IA: costo proyectado interno una vez ejecutadas las '
                 f'salidas plan IA (Jorgelis, Iris, Stipp). Punto de partida ${COSTO_X_PED_ACTUAL:,}/ped (actual).'))
add_bullet(doc, (f'OPERADOR BASE ($2.000/ped tarifa) + jefaturas retenidas (Max + Gerardo + '
                 f'Yohana + 50% Gerente Ops) ~${COSTO_JEFATURAS_INTEG*1000/PEDIDOS_ANO:,.0f}/ped sobre volumen actual'))
add_bullet(doc, (f'OPERADOR OPTIMISTA ($1.500/ped tarifa) + mismas jefaturas retenidas'))

esc_b_data = [['Escenario', 'Pedidos', 'In-house (proy)', 'Operador BASE', 'Operador OPT', 'Decisión']]
for e in esc_top:
    if e['cxp_inhouse_27'] < e['cxp_op_27_opt']:
        decision = 'IN-HOUSE'
        color_d = 'AZUL'
    elif e['cxp_op_27_opt'] < e['cxp_inhouse_27']:
        decision = 'Operador OPT'
        color_d = 'VERDE'
    else:
        decision = 'Indistinto'
        color_d = 'GRIS'
    esc_b_data.append([
        e['nombre'][:20], f'{e["pedidos"]/1000:,.0f}K',
        f'${e["cxp_inhouse_27"]:,.0f}',
        f'${e["cxp_op_27_base"]:,.0f}',
        f'${e["cxp_op_27_opt"]:,.0f}',
        decision,
    ])
add_table(doc, esc_b_data[1:], headers=esc_b_data[0], col_widths=[1.8, 0.8, 1.2, 1.2, 1.2, 1.3])

add_image(doc, os.path.join(CHARTS_DIR, '15_cxp_2027.png'), width=6.3)
add_para(doc, 'Figura 10. Costo por pedido 2027 por escenario — in-house vs operador (base y optimista)',
         italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, 'C.3 Punto de equilibrio — ¿a qué volumen empata el cambio?', level=3)
add_para(doc, ('La pregunta complementaria al análisis por escenario es: independiente del escenario '
               'de venta, ¿a qué volumen anual de pedidos el costo total in-house se iguala al '
               'costo total del operador externo? Por debajo de ese volumen, el operador es más '
               'barato; por encima, el in-house gana eficiencia por absorción de costos fijos.'),
         align='justify')

eq_simple_data = [
    ['Tarifa operador (todo-incl)', 'Modelo operador (cobertura)', 'Volumen de equilibrio', 'Lectura'],
    [f'${BENCH_OP_OPT:,}/ped (OPTIMISTA)', f'+ jefaturas retenidas (${COSTO_JEFATURAS_INTEG/1000:,.0f} MM/año)',
     '~80-120K ped/año',
     'Operador gana SIEMPRE en rango UnionX (180-225K). Cambio justificado financieramente.'],
    [f'${BENCH_OP_BASE:,}/ped (BASE)', f'+ jefaturas retenidas (${COSTO_JEFATURAS_INTEG/1000:,.0f} MM/año)',
     '~250-300K ped/año',
     'Volumen actual 180K y proyección 2027 ~195-225K están BAJO equilibrio. Operador y in-house empatan o gana operador marginalmente. In-house gana por sobre 300K (no realista en 2027).'],
    [f'${BENCH_OP_CONS:,}/ped (CONSERVADOR)', f'+ jefaturas retenidas (${COSTO_JEFATURAS_INTEG/1000:,.0f} MM/año)',
     '~120-150K ped/año',
     'In-house ya gana en rango UnionX. Operador no se justifica.'],
]
add_table(doc, eq_simple_data[1:], headers=eq_simple_data[0], col_widths=[1.6, 1.8, 1.2, 2.1])

add_image(doc, os.path.join(CHARTS_DIR, '14_punto_equilibrio.png'), width=6.3)
add_para(doc, 'Figura 11. Curvas costo total — in-house vs operador (3 tarifas, 3 niveles de eficiencia in-house)',
         italic=True, size=9, color=GRIS_MED, align='center')

add_para(doc, ('Lectura combinando análisis por escenario + punto de equilibrio:'), bold=True, size=10, color=GRIS_OSC)
add_bullet(doc, (f'Volumen actual UnionX: {PEDIDOS_ANO/1000:.0f}K pedidos/año. Proyección 2027 entre '
                 f'{esc_top[1]["pedidos"]/1000:.0f}K (Realista) y {esc_top[4]["pedidos"]/1000:.0f}K (Ambicioso).'))
add_bullet(doc, ('Bajo cualquier escenario realista (incluido el más ambicioso), no superamos los '
                 '300K pedidos/año. Esto significa: para que el operador BASE ($2.000) supere al '
                 'in-house necesitaríamos crecer hacia +60% del volumen actual — no factible en '
                 'horizonte 2027.'))
add_bullet(doc, ('Conclusión combinada: bajo tarifas razonables de mercado (BASE-CONSERVADOR), '
                 'el in-house con plan IA mantiene la ventaja en todo el rango proyectado. Sólo '
                 'la tarifa OPTIMISTA ($1.500) cambia el análisis y justifica una transición.'))

add_heading(doc, 'C.4 Conclusión directa', level=3)
add_para(doc, ('Con el costo retenido in-house elevado por el mantenimiento del Gerente de '
               'Operaciones (50% asignado) + 3 jefaturas operativas integradas, el operador '
               'EXTERNO sólo es competitivo bajo tarifa OPTIMISTA ($1.500/ped o menor). En '
               'cualquier otro escenario, MANTENER OPERACIÓN IN-HOUSE con plan IA aplicado es '
               'la decisión correcta.'), align='justify', bold=True)

add_bullet(doc, ('Volumen actual UnionX: 180K pedidos/año. Costo proyectado in-house con plan '
                 f'IA: ~${COSTO_X_PED_POST_SALIDAS:,.0f}/ped al cierre 2026.'))
add_bullet(doc, ('A volúmenes superiores (escenarios Mejorado/Optimista/Ambicioso 2027), '
                 'in-house gana más eficiencia por mejor absorción de costos fijos.'))
add_bullet(doc, ('El operador externo se vuelve atractivo SÓLO si: (a) se obtiene una tarifa <'
                 ' $1.700/ped vía proceso competitivo; o (b) se descarta el mantenimiento del '
                 'Gerente Ops como parte del modelo retenido.'))
add_bullet(doc, ('Recomendación: mantener operación in-house en 2026-2027 con plan IA. '
                 'Reevaluar trimestralmente si el mercado de operadores logísticos baja sus '
                 'tarifas hacia el rango optimista.'))

doc.add_paragraph()
add_callout(doc, 'RESPUESTA DIRECTA: con el modelo retenido actual (jefaturas operativas + 50% Gerente '
                  'Ops integrados), el operador externo SÓLO conviene si su tarifa todo-incluida es < $1.700/pedido. '
                  'A tarifas $2.000+ pierde frente a in-house con plan IA aplicado en cualquier escenario de venta 2027.',
            color=VERDE)

add_heading(doc, '6.3 Recomendación operador logístico', level=2)
add_bullet(doc, ('2026: NO cambiar a operador en agosto bajo tarifa base ($2.000+). El operador '
                 'optimista ($1.500) sí justifica el cambio, pero requiere negociación contractual '
                 'concreta antes de comprometerse.'))
add_bullet(doc, ('2027: monitorear trimestralmente. Si la tarifa de operador baja a rango $1.500-'
                 '1.800/ped (consolidación del mercado, mayor competencia), evaluar cambio.'))
add_bullet(doc, ('Largo plazo: invertir en automatización in-house (slotting, picking dirigido, '
                 'rotación) para bajar V_in y mover el punto de equilibrio hacia menor volumen, '
                 'haciendo in-house más competitivo.'))
page_break(doc)

# ===== 7. MATRIZ EXTENDIDA ESCENARIOS =====
add_heading(doc, '7. MATRIZ EXTENDIDA DE ESCENARIOS — MC 27% → 33%', level=1)
add_para(doc, ('Para cada nivel de MC% (en pasos de 1pp desde 27% hasta 33%) calculamos la venta '
               'requerida para EBITDA 10%, asumiendo eficiencia GAV gradual y política '
               'inventario gradual coherente con el nivel de palancas. Esto permite ver '
               'múltiples combinaciones realistas.'), align='justify')

add_heading(doc, '7.1 Matriz MC% × Crecimiento V requerido', level=2)
mat_data = [['MC %', 'Ef GAV', 'Pol Inventario', 'Venta req. (MM)', 'Crec V vs 2026', 'UN resultante (MM)', 'Realismo']]
for e in matriz_esc:
    if e['venta'] is None: continue
    realismo = ('Bajo histórico' if e['crec'] <= CREC_25_26*100 else
                'Realista' if e['crec'] <= 15 else
                'Stretch' if e['crec'] <= 25 else 'Aspiracional')
    mat_data.append([
        f'{e["mc"]*100:.0f}%', f'{e["ef_gav"]*100:.1f}%', f'{e["pol_inv"]:.1f} m',
        f'${e["venta"]/1000:,.0f}', f'+{e["crec"]:.0f}%',
        f'${e["un"]/1000:,.0f}', realismo,
    ])
add_table(doc, mat_data[1:], headers=mat_data[0], col_widths=[0.7, 0.8, 1.1, 1.2, 1.2, 1.2, 1.2])

add_image(doc, os.path.join(CHARTS_DIR, '04_matriz_mc.png'), width=6.5)
add_para(doc, 'Figura 11. Matriz extendida — MC% × Crecimiento V requerido para EBITDA 10%',
         italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '7.2 5 Escenarios Top — comparación', level=2)
esc_def = [['Escenario', 'Crec V', 'MC%', 'Ef GAV', 'Pol Inv', 'Venta', 'EBITDA %', 'UN', 'D/E']]
for e in esc_top:
    esc_def.append([
        e['nombre'][:25], f'+{e["crec"]*100:.0f}%', f'{e["mc"]*100:.0f}%',
        f'{e["ef_gav"]*100:.0f}%', f'{e["pol_inv"]:.1f}m',
        f'${e["venta"]/1000:,.0f}', f'{e["ebitda_pct"]*100:.1f}%',
        f'${e["un"]/1000:,.0f}', f'{e["de"]:.2f}x',
    ])
add_table(doc, esc_def[1:], headers=esc_def[0], col_widths=[1.7, 0.7, 0.6, 0.6, 0.6, 0.9, 0.8, 0.9, 0.6])

add_image(doc, os.path.join(CHARTS_DIR, '05_escenarios_top.png'), width=6.3)
add_para(doc, 'Figura 12. EBITDA % resultante por escenario',
         italic=True, size=9, color=GRIS_MED, align='center')

add_image(doc, os.path.join(CHARTS_DIR, '06_cumplimiento.png'), width=6.3)
add_para(doc, 'Figura 13. Matriz cumplimiento — escenarios vs restricciones CEO',
         italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '7.3 Segunda partida: efecto del operador logístico en cada escenario', level=2)
add_para(doc, ('Para cada uno de los 5 escenarios anteriores, evaluamos el efecto financiero de '
               'mantener la operación in-house (default) vs trasladar la operación al operador '
               'logístico externo bajo tarifa base ($2.000) o tarifa optimista ($1.500). El '
               'cambio operador modifica el costo operativo y se traduce directamente en EBITDA '
               'y Utilidad Neta.'), align='justify')

add_heading(doc, 'A. Costo operativo total por escenario × estructura', level=3)
ops_data = [['Escenario', 'Venta (MM)', 'Pedidos', 'Costo OP in-house', 'Costo OP base ($2k)', 'Costo OP opt ($1.5k)', 'Δ OP base', 'Δ OP opt']]
for e in esc_top:
    ops_data.append([
        e['nombre'][:18], f'${e["venta"]/1000:,.0f}', f'{e["pedidos"]/1000:,.0f}K',
        f'${e["costo_op_inh"]/1000:,.0f} MM',
        f'${e["costo_op_base_27"]/1000:,.0f} MM',
        f'${e["costo_op_opt_27"]/1000:,.0f} MM',
        f'${e["delta_op_base"]/1000:+,.0f}',
        f'${e["delta_op_opt"]/1000:+,.0f}',
    ])
add_table(doc, ops_data[1:], headers=ops_data[0], col_widths=[1.5, 0.9, 0.7, 1.1, 1.1, 1.1, 0.8, 0.8])

add_para(doc, ('Lectura: Δ positivo significa que el operador AHORRA respecto a in-house. Δ '
               'negativo significa que el operador es más caro.'), align='justify', italic=True, size=9)

add_heading(doc, 'B. EBITDA y Utilidad Neta por escenario × estructura', level=3)
ebitda_data = [['Escenario', 'In-house EBITDA %', 'In-house UN', 'Op. BASE EBITDA %', 'Op. BASE UN', 'Op. OPT EBITDA %', 'Op. OPT UN']]
for e in esc_top:
    ebitda_data.append([
        e['nombre'][:18],
        f'{e["ebitda_pct"]*100:.1f}%',
        f'${e["un"]/1000:,.0f}',
        f'{e["ebitda_pct_op_base"]*100:.1f}%',
        f'${e["un_op_base"]/1000:,.0f}',
        f'{e["ebitda_pct_op_opt"]*100:.1f}%',
        f'${e["un_op_opt"]/1000:,.0f}',
    ])
add_table(doc, ebitda_data[1:], headers=ebitda_data[0], col_widths=[1.4, 1.1, 0.9, 1.1, 0.9, 1.1, 0.9])

add_image(doc, os.path.join(CHARTS_DIR, '16_efecto_operador.png'), width=6.5)
add_para(doc, 'Figura 14. EBITDA % por escenario × estructura operativa (in-house / operador base / operador optimista)',
         italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, 'C. Lectura clave del efecto operador', level=3)
add_bullet(doc, ('Operador OPTIMISTA ($1.500/ped) mejora el EBITDA en todos los escenarios — el '
                 'ahorro operativo es significativo. Bajo este precio, conviene cambiar.'))
add_bullet(doc, ('Operador BASE ($2.000/ped) mejora marginalmente el EBITDA en escenarios bajos '
                 'y medios. La diferencia se diluye en escenarios de mayor crecimiento (in-house '
                 'absorbe mejor sus fijos).'))
add_bullet(doc, ('Mensaje estratégico: la palanca operador es ATRACTIVA si se logra negociar '
                 'tarifa < $1.700/ped. Bajo tarifa base $2.000, el ahorro existe pero no es '
                 'decisivo. Bajo tarifa $2.500+ destruye valor.'))

add_heading(doc, 'D. Combinaciones más relevantes — Cumplimiento CEO', level=3)
add_para(doc, ('Considerando las 5 restricciones del CEO bajo cada escenario × estructura:'),
         align='justify')
comb_data = [['Combinación', 'EBITDA %', '≥10%?', 'UN (MM)', '≥$333?']]
for e in esc_top:
    for variante, eb, un, c_eb, c_un in [
        ('in-house', e['ebitda_pct']*100, e['un'], e['cumple_ebitda'], e['cumple_un']),
        ('+ Op. BASE', e['ebitda_pct_op_base']*100, e['un_op_base'], e['cumple_ebitda_op_base'], e['cumple_un_op_base']),
        ('+ Op. OPT', e['ebitda_pct_op_opt']*100, e['un_op_opt'], e['cumple_ebitda_op_opt'], e['cumple_un_op_opt']),
    ]:
        comb_data.append([
            f'{e["nombre"][:14]} | {variante}',
            f'{eb:.1f}%', '✓' if c_eb else '✗',
            f'${un/1000:,.0f}', '✓' if c_un else '✗',
        ])
add_table(doc, comb_data[1:], headers=comb_data[0], col_widths=[2.8, 1.0, 0.7, 1.2, 0.7])

add_para(doc, ('Conclusión: el escenario Realista + Operador Optimista ya cumple EBITDA 10% '
               'sin requerir mejoras agresivas en MC. Esto sugiere que el cambio a operador '
               '(si se obtiene tarifa $1.500-1.700) es una palanca tan poderosa como mejorar '
               'el MC del 29% al 33%.'), align='justify', bold=True)

add_heading(doc, '7.4 Factibilidad y recomendación', level=2)
fact = [
    ['Escenario', 'Probabilidad', 'Riesgos'],
    ['Status quo', 'Status quo da pérdida', 'INACEPTABLE — UN negativa.'],
    ['Realista (continuidad)', '60%', f'Crecimiento histórico realista, pero MC 29% no alcanza para EBITDA 10%. UN ${esc_top[1]["un"]/1000:,.0f} MM.'],
    ['Mejorado (palancas moderadas)', '40%', 'Requiere ejecución plan IA + 5 palancas comerciales con éxito moderado.'],
    ['Optimista (palancas plenas)', '20%', 'Requiere mejora compras + USD forward + B2B fuerte + corte canales tóxicos.'],
    ['Ambicioso (stretch)', '8%', 'Coordinación perfecta multi-área + entorno favorable.'],
]
add_table(doc, fact[1:], headers=fact[0], col_widths=[2.0, 1.2, 3.3])

add_para(doc, ('Recomendación: definir como meta corporativa el escenario Mejorado, con plan '
               'de stretch hacia Optimista si Q1 2027 muestra tracción comercial. El '
               'escenario Realista no cumple los objetivos del CEO — debe quedar como piso, '
               'no como meta.'), align='justify', bold=True)
page_break(doc)

# ===== 8. SENSIBILIDAD ±20% =====
add_heading(doc, '8. SENSIBILIDAD ±20% SOBRE 2026', level=1)
add_para(doc, ('Si el FCST 2026 no se cumple, el efecto sobre el patrimonio inicio 2027 es '
               'relevante por el apalancamiento operativo.'), align='justify')
sens_data = [['Δ vs FCST', 'Venta 2026', 'UN 2026', 'Patrimonio inicio 2027', 'D/E inicial', 'Implicancia']]
for s in sens_2026:
    impl = ('Activar Capa 1 contingencia. UN 2027 debe ser ≥$400 MM' if s['delta']<-0.10 else
            'Alerta amarilla. Plan IA + ef GAV prioritarios' if s['delta']<0 else
            'Plan base vigente' if s['delta']==0 else
            'Liberar recursos para inclusiones' if s['delta']<0.15 else
            'Patrimonio robusto, evaluar estrategias ambiciosas')
    sens_data.append([f'{s["delta"]*100:+.0f}%', f'${s["venta26"]/1000:,.0f}',
                       f'${s["un26"]/1000:,.0f}', f'${s["patrim_27"]/1000:,.0f}',
                       f'{s["de_27"]:.2f}x', impl])
add_table(doc, sens_data[1:], headers=sens_data[0], col_widths=[0.8, 1.0, 1.0, 1.3, 0.8, 2.3])
add_image(doc, os.path.join(CHARTS_DIR, '07_sens_2026.png'), width=6.3)
add_para(doc, 'Figura 14. Sensibilidad ±20% en venta 2026', italic=True, size=9, color=GRIS_MED, align='center')

add_para(doc, ('Apalancamiento operativo significativo: -20% en venta convierte la UN +$40 MM '
               f'en pérdida ~-${abs(sens_2026[0]["un26"])/1000:,.0f} MM. +20% multiplicaría la '
               f'UN a ~${sens_2026[-1]["un26"]/1000:,.0f} MM.'), align='justify')
page_break(doc)

# ===== 9. PLAN CONTINGENCIA =====
add_heading(doc, '9. PLAN DE CONTINGENCIA: AJUSTES DESDE LO LEJANO AL CORE', level=1)
add_para(doc, ('Protocolo de ajustes en capas concéntricas desde lo discrecional (capa 1) hasta '
               'el core estratégico (capa 6 inviolable). Las jefaturas operativas no son '
               'inviolables — están en Capa 5 (estructura ajustable vía cambio a operador).'), align='justify')

add_image(doc, os.path.join(CHARTS_DIR, '08_contingencia.png'), width=6.3)
add_para(doc, 'Figura 15. Capas del plan de contingencia', italic=True, size=9, color=GRIS_MED, align='center')

add_heading(doc, '9.1 Detalle por capa', level=2)
for capa in capas_cont:
    add_heading(doc, f'Capa {capa["capa"]}: {capa["nombre"]}', level=3)
    add_para(doc, f'Capacidad recorte anual: ${capa["monto"]/1000:,.0f} MM · Velocidad: {capa["vel"]}',
             italic=True, size=9, color=GRIS_MED)
    for item in capa['items']: add_bullet(doc, item)

add_heading(doc, '9.2 Triggers trimestrales', level=2)
trig_data = [
    ['Trigger', 'Acción', 'Velocidad'],
    ['Q1 < 85% plan', 'Activar Capa 1', 'Mismo Q2'],
    ['Q2 acum < 85% plan', 'Activar Capa 2 + revisión Capa 3', 'Q3'],
    ['Q3 < 90% (ajustado estacionalidad)', 'Capa 3 + acelerar plan IA', 'Q3-Q4'],
    ['Cierre anual < 90% plan', 'Capa 4 + revisión Capa 5 (operador)', 'Año siguiente'],
    ['D/E > 3,5 dos trimestres', 'Plan emergencia: postergar retiros, refinanciar', 'Inmediato'],
]
add_table(doc, trig_data[1:], headers=trig_data[0], col_widths=[2.5, 2.7, 1.3])

total_1a3 = sum(c['monto'] for c in capas_cont[:3])
total_1a5 = sum(c['monto'] for c in capas_cont[:5])
add_heading(doc, '9.3 Capacidad total de defensa', level=2)
def_data = [
    ['Nivel', 'Monto anual (MM)', 'Comentario'],
    ['Capas 1-2', f'${(capas_cont[0]["monto"]+capas_cont[1]["monto"])/1000:,.0f}', 'Sin tocar plantilla'],
    ['Capas 1-3', f'${total_1a3/1000:,.0f}', 'Sin tocar estructura operativa'],
    ['Capas 1-5', f'${total_1a5/1000:,.0f}', 'Preservando core estratégico'],
]
add_table(doc, def_data[1:], headers=def_data[0], col_widths=[2.0, 1.5, 3.0])
page_break(doc)

# ===== 10. RUTA ESTRATÉGICA MULTIFACTOR =====
add_heading(doc, '10. RUTA ESTRATÉGICA MULTIFACTOR', level=1)
add_para(doc, ('Conclusión central del análisis: el cumplimiento del objetivo EBITDA 10% no se '
               'alcanza con una sola palanca. La ruta estratégica es una sumatoria de factores '
               'coordinados, cada uno con su factibilidad, esfuerzo y plazo. Esta sección '
               'descompone el EBITDA gap entre el cierre 2026 y el target 2027, y propone una '
               'ruta combinada realista.'), align='justify')

add_heading(doc, '10.1 Definición del EBITDA gap a cubrir', level=2)
gap_data = [
    ['Concepto', 'Valor (MM CLP)'],
    ['EBITDA 2026 FCST', f'${EBITDA26/1000:,.0f}'],
    ['Venta target 2027 (continuidad histórica +' + f'{CREC_25_26*100:.0f}%)', f'${V_TARGET/1000:,.0f}'],
    ['EBITDA target 2027 (10% × Venta)', f'${EBITDA_TARGET/1000:,.0f}'],
    ['GAP EBITDA a cubrir', f'+${EBITDA_GAP/1000:,.0f}'],
]
add_table(doc, gap_data[1:], headers=gap_data[0], col_widths=[4.5, 2.0])

add_para(doc, (f'El EBITDA actual ($268 MM) representa 4,0% de la venta. Para llegar al 10% '
               'sobre una venta de continuidad ($' + f'{V_TARGET/1000:,.0f}' + ' MM), se necesita '
               'sumar ' + f'~${EBITDA_GAP/1000:,.0f}' + ' MM de EBITDA adicional vía '
               'palancas combinadas.'), align='justify')

add_heading(doc, '10.2 Catálogo de palancas — honesto sobre supuestos y overlaps', level=2)
add_para(doc, ('Reconocimiento crítico: muchas palancas listadas en análisis financieros tienen '
               'doble conteo o asumen mejoras sin base concreta. Esta sección distingue palancas '
               'con BASE CONCRETA (cifras del P&L, contratos firmes) de palancas con SUPUESTO '
               'BLANDO (mejoras de productividad estimadas sin medición).'), align='justify')

add_callout(doc, 'ALERTA METODOLÓGICA: en este catálogo se evita el doble conteo. Por ejemplo, '
                  '"Plan IA salidas" y "eficiencia GAV" no se suman — son la misma cosa visto desde '
                  '2 ángulos. Igual para "cambio operador" y "eficiencia in-house operativa" — son '
                  'rutas ALTERNATIVAS, no aditivas. La ruta combinada de 10.4 selecciona palancas '
                  'NO solapadas.',
            color=NARANJA)

add_heading(doc, 'A. Palancas con BASE CONCRETA (cifras P&L o contratos verificables)', level=3)
base_data = [
    ['Palanca', 'Aporte estimado', 'Origen del cálculo'],
    ['Plan IA salidas (5 cargos confirmados)', '$50-60 MM/año',
     '5 sueldos × cargas patronales 12%, con fechas concretas en plan IA'],
    ['Eliminación Multivende (jun 2026)', '$1,98 MM/año',
     'Costo P&L 2026 Multivende: $1,98 MM. Reemplazado por integración directa'],
    ['Reducción usuarios Odoo (35 → 15)', '$5,5 MM/año',
     'Costo P&L 2026 Odoo: $9,57 MM × 20/35 reducción'],
    ['Crecimiento histórico +8% venta', '$35 MM EBITDA',
     'Continuidad ritmo 2025→2026 (+8%) × margen marginal 24%'],
    ['Política inventario 5,5 → 4 m', '$17 MM/año en GNO (no EBITDA)',
     '1,5 meses × CV/12 × tasa interés 10,5%'],
    ['Subtotal palancas concretas', '~$110 MM/año EBITDA + $17 MM GNO', ''],
]
add_table(doc, base_data[1:], headers=base_data[0], col_widths=[2.5, 1.3, 3.0])

add_heading(doc, 'B. Palancas comerciales (estimadas, requieren ejecución)', level=3)
com_data = [
    ['Palanca comercial', 'Impacto MC (pp)', 'EBITDA equivalente'],
    ['Mejora compras + USD forward', '+0,8-1,5 pp', '$53-100 MM'],
    ['Renegociar tarifas Flex ML+Fala', '+0,3-0,7 pp', '$20-47 MM'],
    ['Marketing digital eficiente', '+0,5-1,0 pp', '$33-67 MM'],
    ['Renegociar comisiones medios de pago web', '+0,2-0,4 pp', '$13-27 MM'],
    ['Crecer share canales propios + B2B', '+1,0-2,5 pp', '$67-167 MM'],
    ['Suma todas (escenario realista)', '+3,0-4,5 pp MC', '$200-300 MM'],
]
add_table(doc, com_data[1:], headers=com_data[0], col_widths=[3.0, 1.5, 2.0])

add_para(doc, ('Nota: estas palancas DEPENDEN de ejecución comercial y negociación. La cifra '
               'realista es lo que la dirección compromete; lo optimista requiere condiciones '
               'favorables que no controlamos plenamente.'), align='justify', italic=True, size=9)

add_heading(doc, 'C. Auditoría GAV — partidas por cuenta del P&L 2026', level=3)
add_para(doc, ('Análisis cuenta por cuenta del GAV 2026 para identificar oportunidades de '
               'eficiencia MÁS ALLÁ del plan IA. No todas las cuentas tienen potencial igual.'),
         align='justify')

gav_audit = [
    ['Cuenta P&L 2026', 'Costo anual (MM)', '% GAV', 'Potencial eficiencia', 'Cómo'],
    ['Sueldos GAV', '$1.180', '76,2%', '5-7% (~$60-80 MM)',
     'Plan IA salidas + audit consolidación duplicaciones (TradeMkt, KAMs pequeños, asesorías ext)'],
    ['Oficina y Arriendos', '$201', '13,0%', '5-10% (~$10-20 MM)',
     'Renegociar arriendos al vencimiento · evaluar reducción m² oficina con teletrabajo parcial'],
    ['Marketing Branding', '$43', '2,8%', 'Reasignación, no recorte',
     'Auditoría ROI por campaña · reducir branding genérico, aumentar SEO/contenido orgánico'],
    ['Suscripciones', '$31', '2,0%', '20-30% (~$6-10 MM)',
     'Plan IA confirmado: -Multivende, -Odoo 35→15. Auditar SaaS no esenciales'],
    ['Movilización/Transporte', '$23', '1,5%', '15-20% (~$4 MM)',
     'Revisar política viajes · alternativas digitales · reembolso vs corporativo'],
    ['Depreciación', '$18', '1,2%', '0% (técnico)',
     'No ajustable. Sólo cambia con nuevas inversiones (PP&E)'],
    ['Seguros', '$17', '1,1%', '10-15% (~$2 MM)',
     'Renegociar pólizas en bloque · revisar coberturas redundantes'],
    ['I+D Productos', '$17', '1,1%', 'Mantener o priorizar',
     'No recortar si genera SKUs con buen MC. Priorizar I+D rentable'],
    ['Capacitación personal', '$6', '0,4%', 'Mantener',
     'Inversión en capacidad. Recortable solo en escenario contingencia'],
    ['Asesorías varias', '$4', '0,3%', '30-50% (~$1-2 MM)',
     'Revisar contratos asesores · consolidar con plan IA reduciendo necesidad externa'],
    ['Impuestos/Patentes', '$4', '0,3%', '0%', 'No ajustable'],
    ['Asesoría contable externa', '$2', '0,1%', '50% (~$1 MM)',
     'Con plan IA LCV + 1 analista in-house, evaluar reducir asesoría externa'],
    ['TOTAL GAV', '$1.548', '100%', '~$85-125 MM eficiencia potencial', '5,5-8% del GAV total'],
]
add_table(doc, gav_audit[1:], headers=gav_audit[0], col_widths=[1.7, 0.9, 0.5, 1.3, 2.4])

add_callout(doc, 'AUDITORÍA GAV: el potencial total de eficiencia GAV es ~$85-125 MM/año (5-8%). '
                  'De ese total, el plan IA cubre ~$60 MM. El resto ($25-65 MM) requiere acciones '
                  'puntuales en arriendos, suscripciones, viajes, seguros y asesorías. Ninguna por sí '
                  'sola es transformacional, pero sumadas aportan a la ruta multifactor.',
            color=AZUL_OSC)

add_heading(doc, 'D. Palancas con SUPUESTO BLANDO (requieren validación)', level=3)
blando_data = [
    ['Palanca', 'Aporte hipotético', 'Por qué es blando'],
    ['Eficiencia in-house operativa (slotting/picking)', '$30-60 MM/año',
     'Asume mejora productividad sin medición previa. No hay benchmark in-house ni '
     'baseline cuantitativo de productividad actual'],
    ['Reducción mermas/devoluciones', '$10-20 MM/año',
     'Asume mejora plan IA SAC. Sin base de mermas reales mensuales no es cuantificable'],
    ['Refinanciar deuda a tasa más baja', '$30-50 MM/año',
     'Asume condiciones de mercado mejores. Depende de ratings, garantías, contexto macro'],
    ['Optimización tributaria', '$10-30 MM/año',
     'Sin estudio tributario específico no es cuantificable'],
]
add_table(doc, blando_data[1:], headers=blando_data[0], col_widths=[2.5, 1.5, 2.8])

add_para(doc, ('Recomendación: las palancas blandas deben validarse con estudios específicos '
               'antes de comprometerlas en el plan 2027. Sirven como upside potencial, no como '
               'línea base.'), align='justify', italic=True, size=9)

add_heading(doc, '10.3 Factibilidad por palanca clave', level=2)
add_para(doc, ('Las palancas se ordenan por relación impacto / esfuerzo (las más altas en ambas '
               'son prioritarias).'), align='justify')

fact_data = [
    ['Palanca', 'Aporte realista', 'Esfuerzo', 'Plazo', 'Quick win?'],
    ['Crecimiento histórico +8% venta', '$35 MM', 'Medio', 'Continuo', 'No · ya en ejecución'],
    ['Mejora compras + USD forward (MC +1pp)', '$67 MM', 'Medio', 'H1 2027', 'Sí · negociable corto plazo'],
    ['Plan IA salidas pleno (6 cargos)', '$60 MM', 'Bajo', 'jul-oct 2026', 'Sí · plan ya armado'],
    ['Eficiencia in-house operativo (slotting)', '$60 MM', 'Medio-Alto', 'H2 2027', 'No · requiere transformación'],
    ['Crecer canales propios + B2B (MC +1pp)', '$67 MM', 'Alto', '12-18 m', 'No · transformación comercial'],
    ['Cambio operador (tarifa BASE)', '$52 MM', 'Alto', 'Q3 2026', 'No · indemnización + transición'],
    ['Marketing digital eficiente (MC +0,7pp)', '$47 MM', 'Medio', 'H1 2027', 'Sí · reasignación gasto'],
    ['Audit GAV (consolidación duplicaciones)', '$25 MM', 'Medio', 'Q1 2027', 'Sí · decisión interna'],
    ['Negociación volumen proveedores (-1% CD)', '$31 MM', 'Medio', 'Q1 2027', 'Sí · renegociación contractual'],
    ['Tarifa Flex ML+Fala (MC +0,5pp)', '$33 MM', 'Bajo-Medio', 'H1 2027', 'Sí · negociación con couriers'],
    ['Política inventario 5,5 → 4 m (KT)', '$15 MM GNO', 'Medio', 'Q1-Q2 2027', 'Parcial · requiere plan slotting'],
    ['Refinanciar deuda a tasa más baja', '$30 MM', 'Alto', '6-12 m', 'No · requiere mercado'],
]
add_table(doc, fact_data[1:], headers=fact_data[0], col_widths=[2.3, 1.0, 0.8, 0.9, 1.5])

add_heading(doc, '10.4 Ruta combinada recomendada', level=2)
add_para(doc, ('Selección de palancas sin doble conteo para alcanzar el target EBITDA con '
               'crecimiento de venta al ritmo histórico (+8%):'), align='justify')

ruta_data = [['Palanca', 'Aporte EBITDA (MM)', 'Categoría']]
for cat, p, val in ruta_recomendada:
    ruta_data.append([p, f'+${val}', cat])
ruta_data.append(['TOTAL RUTA', f'+${TOTAL_RUTA} MM', ''])
ruta_data.append(['EBITDA 2026 base', f'${EBITDA26/1000:,.0f} MM', ''])
ruta_data.append(['EBITDA 2027 resultante', f'${EBITDA_27_RUTA/1000:,.0f} MM', f'{EBITDA_27_RUTA_PCT*100:.1f}%'])
ruta_data.append(['vs EBITDA target 10%', f'${EBITDA_TARGET/1000:,.0f} MM', '10,0%'])
add_table(doc, ruta_data[1:], headers=ruta_data[0], col_widths=[3.5, 1.5, 1.5])

add_image(doc, os.path.join(CHARTS_DIR, '17_waterfall_palancas.png'), width=6.5)
add_para(doc, 'Figura 16. Waterfall EBITDA 2026 → 2027 — descomposición de aporte por palanca',
         italic=True, size=9, color=GRIS_MED, align='center')

add_para(doc, ('Resultado: EBITDA 2027 alcanza ' + f'{EBITDA_27_RUTA_PCT*100:.1f}%' +
               ' sobre venta. ' +
               ('Cumple el objetivo CEO 10% con holgura.' if EBITDA_27_RUTA_PCT >= 0.10
                else f'Cubre {(EBITDA_27_RUTA/EBITDA_TARGET)*100:.0f}% del objetivo CEO. Falta cerrar el gap residual con palancas adicionales o ajustando palancas hacia escenario optimista.')),
         align='justify', bold=True)

add_para(doc, ('Variante: si la palanca de eficiencia in-house (slotting/picking) no se ejecuta, '
               'se puede sustituir por el cambio a operador externo bajo tarifa optimista '
               '($1.500/ped). El aporte resultante es similar pero con mayor riesgo de ejecución '
               'contractual.'), align='justify', italic=True, size=9)

add_heading(doc, '10.6 Factibilidad realista — cumplir TODOS los objetivos vs por prioridad', level=2)
add_para(doc, ('Conclusión crítica del análisis: NO es factible cumplir simultáneamente los 5 '
               'objetivos del CEO con palancas de continuidad. La combinación que cumple los 5 '
               'objetivos (incluido EBITDA 10% completo) requiere ejecución cercana al escenario '
               'optimista en todas las palancas — riesgo de ejecución alto. Por eso este '
               'documento presenta el cumplimiento por prioridad descendente: cuánto se logra '
               'con qué nivel de ejecución.'), align='justify')

# Tabla de cumplimiento por prioridad
add_heading(doc, 'Orden de prioridad de los objetivos del CEO', level=3)
prior_data = [
    ['Prioridad', 'Objetivo', 'Criterio cuantitativo', 'Difícil/Fácil de cumplir'],
    ['1ª (más alta)', 'D/E en rango bajo riesgo (2,5-3,5)', f'Patrimonio post-retiros entre $580-$800 MM', 'Fácil — UN positiva basta'],
    ['2ª', 'Sueldo CEO crece 10-15%', f'Ajuste salarial general 12% ya aplicado', 'Confirmado en supuestos'],
    ['3ª', 'Sueldo Erich vitalicio', 'Mantener $5,21 MM/m', 'Confirmado'],
    ['4ª', 'Retirar $100 MM (UN ≥ $333 MM)', 'UN 2027 ≥ $333 MM', 'Medio — escenario realista llega a ~$200-250 MM'],
    ['5ª (más alta exigencia)', 'EBITDA 10%', f'EBITDA 2027 ≥ ${EBITDA_TARGET/1000:,.0f} MM', 'Difícil — requiere combinación palancas multifactor'],
]
add_table(doc, prior_data[1:], headers=prior_data[0], col_widths=[1.4, 2.0, 1.9, 1.4])

add_heading(doc, 'Escenarios por nivel de cumplimiento', level=3)
add_para(doc, ('Tabla que muestra, según el nivel de ejecución alcanzado, qué objetivos se '
               'cumplen (de mayor a menor prioridad):'), align='justify')

niv_data = [
    ['Nivel ejecución', 'Palancas activadas', 'EBITDA % logrado', 'UN logrado', 'Objetivos cumplidos'],
    ['Status quo', 'Ninguna palanca · sin cambios',
     '< 4%', 'Pérdida potencial',
     'Sólo 2ª y 3ª (sueldos). Falla 1ª, 4ª y 5ª. INACEPTABLE.'],
    ['Mínimo viable', 'Crecimiento histórico +8% · Plan IA salidas pleno',
     '~5-6%', '~$150-200 MM',
     'Cumple 1ª, 2ª, 3ª. Falla 4ª y 5ª.'],
    ['Realista', 'Anterior + MC 29% (compras + mkt eficiente) + ef GAV 5%',
     '~7-8%', '~$280-330 MM',
     'Cumple 1ª, 2ª, 3ª, cerca de 4ª. Falla 5ª.'],
    ['Mejorado', 'Anterior + MC 31% + ef GAV 8% + neg proveedores',
     '~9-10%', '~$420-500 MM',
     'Cumple 1ª, 2ª, 3ª, 4ª. Cerca de 5ª.'],
    ['Completo (todos los objetivos)', 'Mejorado + 5ª palanca MC (canales propios + B2B) + cambio operador OPT o eficiencia in-house plena',
     '~10-12%', '~$500-650 MM',
     'Cumple TODOS los 5 objetivos del CEO.'],
]
add_table(doc, niv_data[1:], headers=niv_data[0], col_widths=[1.5, 2.6, 0.9, 0.9, 1.5])

add_heading(doc, 'Lectura honesta para el directorio', level=3)
add_bullet(doc, ('Objetivos 1, 2, 3 (D/E, sueldos CEO y Erich): se cumplen con relativa facilidad '
                 'en cualquier escenario donde la UN sea positiva.'))
add_bullet(doc, ('Objetivo 4 (UN ≥ $333 MM para retirar $100 MM): se cumple en el escenario '
                 'Mejorado y superiores. En el Realista quedamos al borde (~$300 MM).'))
add_bullet(doc, ('Objetivo 5 (EBITDA 10%): es el más exigente. Requiere combinación multifactor '
                 'con MC ≥ 31% Y eficiencia GAV ≥ 8% Y/O cambio operador atractivo. Si solo '
                 'se ejecutan 2 palancas, el EBITDA queda en 7-9%, no en 10%.'))
add_bullet(doc, ('Honesto al directorio: comprometer EBITDA 10% es comprometer un escenario '
                 'Mejorado-Optimista. Es factible pero exigente. Comprometer UN $333 MM es '
                 'factible con palancas moderadas. Comprometer D/E y sueldos es trivial.'))

add_heading(doc, '10.7 Pilares de la ruta estratégica', level=2)
add_para(doc, 'La ruta para acercarse al objetivo CEO se construye sobre 4 pilares:',
         bold=True, size=10, color=AZUL_OSC)
add_bullet(doc, ('Pilar 1 — Venta: crecer al ritmo histórico (+8%). NO requiere apostar a un '
                 'crecimiento agresivo de mercado. Es lo que la empresa demostró que puede hacer '
                 '2025→2026.'))
add_bullet(doc, ('Pilar 2 — Márgenes: mejorar MC de 27% a 31% combinando 3 palancas '
                 'simultáneas (compras + USD forward + marketing digital eficiente + share '
                 'canales propios). NO depende de renegociar marketplaces.'))
add_bullet(doc, ('Pilar 3 — Eficiencia GAV: ejecutar plan IA al 100% + audit consultora '
                 'consolidando duplicaciones. Reduce GAV ~$85 MM/año en régimen.'))
add_bullet(doc, ('Pilar 4 — Operación: elegir entre eficientar in-house (slotting + picking) '
                 'O cambiar a operador externo bajo tarifa atractiva ($1.500-2.000). Ambas '
                 'rutas son válidas; la decisión depende del benchmark real de mercado.'))

add_para(doc, ('Adicionalmente, las palancas de Capital de Trabajo (política inventario, días '
               'CxC/CxP) y Estructura Financiera (refinanciamiento) NO aportan EBITDA directo '
               'pero reducen GNO y mejoran UN. Son palancas complementarias importantes para '
               'cumplir las restricciones de UN $333 MM y D/E.'), align='justify')

add_heading(doc, '10.8 Próximos pasos', level=2)
add_bullet(doc, 'Q4 2026: validar este análisis con CFO y CEO. Definir Ruta A o B como meta.')
add_bullet(doc, 'Dic 2026: presentar al directorio. Aprobación presupuesto 2027 alineado a la ruta.')
add_bullet(doc, 'Ene 2027: kick-off por pilar: plan comercial · plan IA pleno · plan eficiencia GAV · decisión operador o eficiencia in-house · plan inventario 4m.')
add_bullet(doc, 'Trimestral 2027: comité de seguimiento con KPIs por palanca + triggers contingencia.')

add_para(doc, 'Mensaje final', bold=True, size=11, color=AZUL_OSC)
add_para(doc, ('Los objetivos del CEO se alcanzan combinando palancas de 4 pilares: venta '
               'histórica + márgenes + GAV + operación. Ningún factor por separado es suficiente, '
               'pero la combinación coordinada cubre el gap de $' + f'{EBITDA_GAP/1000:,.0f}' +
               ' MM en EBITDA. La ejecución exige disciplina trimestral y compromisos cruzados '
               'entre comercial, operaciones, finanzas y compras. La rentabilidad no es resultado '
               'de un esfuerzo individual; es el resultado de la coordinación de toda la empresa.'),
         align='justify', italic=True)
page_break(doc)

# ===== 11. ANEXOS =====
add_heading(doc, '11. ANEXOS', level=1)
add_heading(doc, '11.1 Suscripciones SaaS 2026 vs 2027', level=2)
saas = [
    ['Concepto', '2026 (MM)', '2027 (MM)', 'Comentario'],
    ['Yuju', '$7,36', '$7,36', 'Se mantiene'],
    ['Multivende', '$1,98', '$0', 'Eliminado desde junio 2026'],
    ['Odoo (35→15)', '$9,57', '$4,10', '-57%'],
    ['Claude API', '~$2,5', '$24,8', 'Plan IA pleno'],
    ['ChatGPT', '$0,98', '$1,02', 'UF'],
    ['Otros', '~$8,5', '~$8,8', 'UF'],
    ['Total', f'${SUSCRIPCIONES26/1000:,.1f}', f'${calc_suscripciones_2027()/1000:,.1f}', ''],
]
add_table(doc, saas[1:], headers=saas[0], col_widths=[1.7, 1.0, 1.0, 2.3])

add_heading(doc, '11.2 Glosario técnico', level=2)
gloss = [
    ['Término', 'Definición'],
    ['MC%', '(Venta - Costo Directo - Otros Costos Explot) / Venta'],
    ['EBITDA', 'EBIT + Depreciación'],
    ['D/E', 'Deuda Financiera / Patrimonio'],
    ['Punto de equilibrio (P*)', 'Volumen donde Costo_in-house = Costo_operador'],
    ['Jefaturas integradas', 'Max Bellolio (Post Venta) + Gerardo Ortega (Logística) + Yohana Grisman (Facturación) + Gerente Operaciones (50% sueldo asignado a Ops) — operan DENTRO del operador externo'],
    ['Tarifa todo-incluida operador', 'Incluye arriendo, sueldos operativos, insumos del operador'],
    ['Costo por pedido in-house actual', '$2.812 (KPI app · promedio 16 meses)'],
]
add_table(doc, gloss[1:], headers=gloss[0], col_widths=[1.8, 4.5])

add_para(doc, '', size=8)
add_para(doc, 'Documento generado: mayo 2026 · Versión 3', italic=True, size=9, color=GRIS_MED, align='center')

doc.save(OUTPUT_DOCX)
print(f"\n✅ Word v3 generado: {OUTPUT_DOCX}")
print(f"   Tamaño: {os.path.getsize(OUTPUT_DOCX)/1024:.0f} KB")
